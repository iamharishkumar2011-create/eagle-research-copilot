"""
EAGLE — Research Copilot for GLIM Gurgaon Faculty
A single-file Streamlit prototype implementing the 5-agent research workflow:
  1. Problematizer  — sharpens fuzzy idea into a research question
  2. Theorist       — builds theoretical lens and construct map
  3. Literature Scout — pulls real papers from OpenAlex, builds lit table
  4. Method Designer — recommends method, generates survey instrument
  5. Reviewer       — devil's advocate critique at target journal level

Run locally:
    pip install streamlit anthropic requests python-docx
    export ANTHROPIC_API_KEY=sk-ant-...
    streamlit run sherpa_app.py

Single-file on purpose so you can hand it to anyone.
"""

import json
import sqlite3
import uuid
import datetime as dt
import io
from typing import Any

import streamlit as st
import requests
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches

# ---------------------------------------------------------------------------
# CONFIG
# ---------------------------------------------------------------------------

MODEL_HEAVY = "claude-3-5-sonnet-latest"   # Problematizer, Theorist, Reviewer
MODEL_LIGHT = "claude-3-5-haiku-latest"    # Lit scout formatting, method items
import os
_is_cloud = os.environ.get("STREAMLIT_SERVER_HEADLESS") or os.environ.get("HOME", "").startswith("/home/appuser")
DB_PATH = "/tmp/eagle.db" if _is_cloud else "eagle.db"
OPENALEX_BASE = "https://api.openalex.org/works"

STAGES = [
    ("stage_1_problematization", "1. Problematize"),
    ("stage_2_theorization",     "2. Theorize"),
    ("stage_3_literature",       "3. Literature"),
    ("stage_4_method",           "4. Method"),
    ("stage_5_review",           "5. Reviewer"),
    ("stage_6_full_paper",       "6. Full Paper"),
]

# ---------------------------------------------------------------------------
# DATABASE
# ---------------------------------------------------------------------------

def db_init():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS sessions (
            session_id TEXT PRIMARY KEY,
            faculty_id TEXT,
            title TEXT,
            canvas_json TEXT,
            created_at TEXT,
            updated_at TEXT
        )
    """)
    conn.commit()
    conn.close()

def db_save(canvas: dict):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        INSERT INTO sessions (session_id, faculty_id, title, canvas_json, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?)
        ON CONFLICT(session_id) DO UPDATE SET
            canvas_json=excluded.canvas_json,
            updated_at=excluded.updated_at,
            title=excluded.title
    """, (
        canvas["session_id"],
        canvas["faculty_id"],
        canvas.get("title", "Untitled"),
        json.dumps(canvas),
        canvas["created_at"],
        dt.datetime.utcnow().isoformat(),
    ))
    conn.commit()
    conn.close()

def db_list(faculty_id: str):
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(
        "SELECT session_id, title, updated_at FROM sessions WHERE faculty_id=? ORDER BY updated_at DESC",
        (faculty_id,)
    ).fetchall()
    conn.close()
    return rows

def db_load(session_id: str) -> dict | None:
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute("SELECT canvas_json FROM sessions WHERE session_id=?", (session_id,)).fetchone()
    conn.close()
    return json.loads(row[0]) if row else None

# ---------------------------------------------------------------------------
# CANVAS — empty scaffold matching the schema
# ---------------------------------------------------------------------------

def new_canvas(faculty_id: str) -> dict:
    now = dt.datetime.utcnow().isoformat()
    return {
        "session_id": str(uuid.uuid4()),
        "faculty_id": faculty_id,
        "title": "Untitled research session",
        "created_at": now,
        "updated_at": now,
        "current_stage": "stage_1_problematization",
        "stage_1_problematization": {
            "fuzzy_idea": "",
            "socratic_exchange": [],
            "candidate_rqs": [],
            "selected_rq": "",
            "faculty_notes": "",
        },
        "stage_2_theorization": {
            "candidate_theories": [],
            "selected_theory": "",
            "construct_map": {},
            "hypotheses": [],
            "faculty_notes": "",
        },
        "stage_3_literature": {
            "search_queries_used": [],
            "lit_table": [],
            "synthesis": "",
            "identified_gap": "",
            "faculty_notes": "",
        },
        "stage_4_method": {
            "research_approach": "quantitative",
            "recommended_method": "",
            "method_rationale": "",
            "sample_plan": {},
            "instrument": {"blocks": [], "attention_checks": [], "demographics": []},
            "interview_protocol": {"sections": [], "probes": [], "sampling_strategy": ""},
            "ethics_flags": [],
            "faculty_notes": "",
        },
        "stage_5_review": {
            "reviewer_critique": [],
            "overall_readiness_score": 0,
            "target_journal": "JAMS",
            "journal_specific_notes": "",
            "faculty_notes": "",
        },
        "stage_6_full_paper": {
            "introduction": "",
            "literature_review": "",
            "hypotheses_text": "",
            "method_section": "",
            "theoretical_implications": [],
            "practical_implications": [],
            "revised_notes": "",
            "generated": False,
        },
    }

# ---------------------------------------------------------------------------
# OPENALEX — real literature retrieval, no hallucination
# ---------------------------------------------------------------------------

def openalex_search(query: str, per_page: int = 25) -> list[dict]:
    """Hit OpenAlex. Returns list of paper dicts. No API key required."""
    params = {
        "search": query,
        "per_page": per_page,
        "filter": "type:article,from_publication_date:2012-01-01",
        "sort": "cited_by_count:desc",
        "mailto": "research@glim.edu.in",  # polite pool
    }
    try:
        r = requests.get(OPENALEX_BASE, params=params, timeout=20)
        r.raise_for_status()
        data = r.json()
    except Exception as e:
        st.warning(f"OpenAlex error for '{query}': {e}")
        return []

    results = []
    for w in data.get("results", []):
        authors = ", ".join(
            a["author"]["display_name"]
            for a in (w.get("authorships") or [])[:4]
        )
        venue = ""
        pl = w.get("primary_location") or {}
        src = pl.get("source") or {}
        venue = src.get("display_name", "") or ""
        results.append({
            "openalex_id": w.get("id", ""),
            "doi": w.get("doi", ""),
            "authors": authors,
            "year": w.get("publication_year", ""),
            "title": w.get("title", ""),
            "venue": venue,
            "abstract": reconstruct_abstract(w.get("abstract_inverted_index")),
            "cited_by": w.get("cited_by_count", 0),
        })
    return results

def reconstruct_abstract(inv_idx: dict | None) -> str:
    """OpenAlex stores abstracts as inverted indexes to avoid copyright issues."""
    if not inv_idx:
        return ""
    positions = []
    for word, idxs in inv_idx.items():
        for i in idxs:
            positions.append((i, word))
    positions.sort()
    return " ".join(w for _, w in positions)

# ---------------------------------------------------------------------------
# CLAUDE CALLS
# ---------------------------------------------------------------------------

def get_client() -> Anthropic:
    api_key = st.session_state.get("anthropic_api_key", "").strip()
    if not api_key:
        try:
            api_key = st.secrets.get("ANTHROPIC_API_KEY", "").strip()
        except:
            pass
    if not api_key:
        api_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        st.error("Please enter your Anthropic API key in the sidebar.")
        st.stop()
    return Anthropic(api_key=api_key)
def claude_call(system: str, user: str, model: str = MODEL_HEAVY, max_tokens: int = 4000, return_meta: bool = False):
    client = get_client()
    try:
        resp = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            system=system,
            messages=[{"role": "user", "content": user}],
        )
        text = resp.content[0].text if resp.content else ""
        if return_meta:
            return text, resp.stop_reason
        return text
    except Exception as e:
        st.error(f"**Anthropic API Error:** {str(e)}")
        st.info(f"Model ID used: {model}")
        st.stop()

def _extract_json(raw: str) -> dict:
    """Try multiple strategies to extract JSON from Claude's response."""
    import re
    cleaned = raw.strip()
    # Strategy 1: direct parse
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
    # Strategy 2: strip code fences
    if "```" in cleaned:
        match = re.search(r"```(?:json)?\s*\n?(.*?)```", cleaned, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(1).strip())
            except json.JSONDecodeError:
                pass
    # Strategy 3: find first { ... last }
    first_brace = cleaned.find("{")
    last_brace = cleaned.rfind("}")
    if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
        try:
            return json.loads(cleaned[first_brace:last_brace + 1])
        except json.JSONDecodeError:
            pass
    raise json.JSONDecodeError("No valid JSON found", cleaned, 0)

def claude_json(system: str, user: str, model: str = MODEL_HEAVY, max_retries: int = 2, max_tokens: int = 4000) -> dict:
    """Call Claude and parse JSON from the response. Retries on failure."""
    last_raw = ""
    for attempt in range(max_retries + 1):
        extra = "\n\nIMPORTANT: Return ONLY valid JSON, no prose, no markdown fences." if attempt > 0 else ""
        raw, stop_reason = claude_call(system, user + extra,
                          model=model, max_tokens=max_tokens, return_meta=True)
        last_raw = raw
        if not raw or not raw.strip():
            if attempt < max_retries:
                st.warning(f"Empty response from Claude (attempt {attempt+1}). Retrying...")
                continue
            st.error("Claude returned an empty response after all retries.")
            raise json.JSONDecodeError("Empty response", "", 0)
        if stop_reason == "max_tokens":
            st.warning(f"Response was truncated (hit token limit). Retrying with simplified prompt...")
            # On retry for truncation, ask for fewer items
            extra = "\n\nCRITICAL: Keep your response SHORT. Limit lit_table to 10 papers max. Return ONLY valid JSON."
            raw, stop_reason = claude_call(system, user + extra,
                              model=model, max_tokens=max_tokens, return_meta=True)
            last_raw = raw
        try:
            return _extract_json(raw)
        except json.JSONDecodeError:
            if attempt < max_retries:
                st.warning(f"JSON parse failed (attempt {attempt+1}/{max_retries+1}). Retrying...")
                continue
            st.error("Couldn't parse JSON from Claude after retries. Raw output:")
            st.code(raw[:2000] if raw else "(empty)")
            raise

# ---------------------------------------------------------------------------
# AGENT PROMPTS (condensed from the full specs — full versions in prompts.md)
# ---------------------------------------------------------------------------

SYS_PROBLEMATIZER_QUESTION = """You are the Problematizer agent in Eagle, a research copilot for
marketing/management faculty at a top Indian business school (GLIM Gurgaon).

You are talking to an experienced researcher. Match that register: intellectually
ambitious, concise, never condescending, no filler. Use Alvesson & Sandberg's
problematization method.

Your job RIGHT NOW: based on the conversation so far, ask ONE sharp Socratic
question that pushes the faculty toward clarity on the phenomenon, tension,
challenged assumption, stakeholders, or what would count as a surprising finding.

Never propose theories, methods, or citations. Never ask more than one question
at a time. Output ONLY the question text, no preamble.
"""

SYS_PROBLEMATIZER_FINALIZE = """You are the Problematizer agent. Based on the fuzzy idea and the Socratic
exchange below, produce THREE candidate research questions. Each must:
- Be empirically answerable
- Name the focal phenomenon
- Imply a tension or contrast
- Avoid generic "impact of X on Y" framing

Return ONLY valid JSON, no prose, no code fences:
{
  "candidate_rqs": [
    {"rq": "...", "tradeoff": "one sentence on what this framing gains/loses"},
    {"rq": "...", "tradeoff": "..."},
    {"rq": "...", "tradeoff": "..."}
  ]
}
"""

SYS_THEORIST = """You are the Theorist agent in Eagle. The faculty has selected a research
question. Your job is to propose theoretical lenses and build a construct map
that a JAMS/JCR/JMR reviewer would find defensible.

Propose 2-3 candidate theories, recommend one, then build a construct map with
focal IV(s), focal DV(s), mediators, moderators, controls, and boundary
conditions. Derive 3-5 hypotheses.

Never fabricate theoretical claims. If unsure, say so. Push back if the
construct map reduces to "A affects B" — a top-tier paper needs mechanism.

Return ONLY valid JSON matching this shape, no prose, no fences:
{
  "candidate_theories": [
    {"name": "...", "core_claim": "...", "strengths": "...", "weaknesses": "...", "seminal_refs": ["...", "..."]}
  ],
  "recommended_theory": "...",
  "recommendation_rationale": "...",
  "construct_map": {
    "focal_iv": [{"name": "...", "definition": "..."}],
    "focal_dv": [{"name": "...", "definition": "..."}],
    "mediators": [{"name": "...", "rationale": "..."}],
    "moderators": [{"name": "...", "rationale": "..."}],
    "controls": ["..."],
    "boundary_conditions": "..."
  },
  "hypotheses": [
    {"id": "H1", "text": "...", "type": "direct|mediation|moderation|moderated_mediation"}
  ]
}
"""

SYS_LITSCOUT_QUERIES = """You are the Literature Scout. Given the construct map and theory below,
generate 5 targeted OpenAlex search queries. Each query should be 3-6 words,
combining constructs, theory, and context. Diversify — don't repeat the same
core phrase.

Return ONLY valid JSON, no fences:
{"queries": ["query 1", "query 2", "query 3", "query 4", "query 5"]}
"""

SYS_LITSCOUT_SYNTHESIZE = """You are the Literature Scout. You have been given a set of REAL papers
retrieved from OpenAlex (author, year, title, venue, abstract, citations).
Your job: (a) score each paper 0-3 on relevance to the focal RQ, (b) build a
structured lit table of the top ~15, (c) write a 350-450 word synthesis that
groups papers into themes, names what the literature collectively establishes,
and identifies 2-3 specific gaps.

HARD RULES:
- Use ONLY the papers provided. Do not add any paper not in the input list.
- Never invent DOIs, authors, or years.
- Never quote more than 10 words from any abstract. Paraphrase findings.
- If the abstract is empty or uninformative, put "requires full-text review"
  in the findings/gap fields.

Return ONLY valid JSON, no fences:
{
  "lit_table": [
    {
      "authors": "...",
      "year": 2020,
      "venue": "...",
      "theory": "...",
      "context": "...",
      "method": "...",
      "key_constructs": "...",
      "findings": "...",
      "gap_identified": "...",
      "relevance_to_rq": "...",
      "doi": "...",
      "openalex_id": "..."
    }
  ],
  "synthesis": "...",
  "identified_gap": "..."
}
"""

SYS_METHOD = """You are the Method Designer agent in Eagle. Given the construct map,
hypotheses, and literature context, recommend a method and generate a survey
instrument.

For each construct, generate 4-5 survey items. Where a well-known validated
scale exists (e.g., Peck & Shu 2009 for psychological ownership, Brakus et al.
2009 for brand experience), cite the source and adapt items. Where none exists,
generate new items following construct-definition-first logic and flag them as
"newly generated — requires pretest."

Include: 2 attention checks, reverse-coded items (at least 20% per scale),
CMB marker (Simmering et al. 2015 blue attitude), and demographics block.

Never pad. A good survey is tight.

Return ONLY valid JSON, no fences:
{
  "recommended_method": "...",
  "method_rationale": "...",
  "sample_plan": {
    "target_n": 0,
    "power_analysis": "...",
    "recruitment": "...",
    "inclusion_criteria": "..."
  },
  "instrument": {
    "blocks": [
      {
        "block_name": "...",
        "construct": "...",
        "scale_source": "...",
        "validated": true,
        "items": [
          {"id": "X1", "text": "...", "scale": "7-point Likert", "reverse_coded": false}
        ]
      }
    ],
    "attention_checks": [{"id": "AC1", "text": "..."}],
    "cmb_marker": {"construct": "...", "items": []},
    "demographics": ["age", "gender", "..."]
  },
  "ethics_flags": ["..."]
}
"""

SYS_REVIEWER = """You are the Reviewer agent in Eagle. You simulate a demanding but fair
reviewer at the target journal. Read the ENTIRE canvas provided. Your job is
to find the concerns that would get this paper desk-rejected or sent for
major revision BEFORE submission.

Evaluate against six criteria:
a) Theoretical contribution
b) Problematization (is the gap real?)
c) Identification and causal inference
d) Measurement (scales, CMB, social desirability)
e) Generalizability and boundary conditions
f) Alternative explanations

Generate 6-10 concerns total. Be direct. Soft feedback wastes the faculty's
time. Never recommend rejection — recommend revision paths. Assign an overall
readiness score 0-10 for the target journal.

Return ONLY valid JSON, no fences:
{
  "reviewer_critique": [
    {
      "concern_id": "R1",
      "category": "Theoretical contribution|Problematization|Identification|Measurement|Generalizability|Alternative explanations",
      "severity": "low|medium|high",
      "issue": "...",
      "suggested_fix": "...",
      "linked_stage": "stage_2_theorization"
    }
  ],
  "overall_readiness_score": 7.0,
  "journal_specific_notes": "..."
}
"""

SYS_METHOD_QUALITATIVE = """You are the Method Designer agent in Eagle. The faculty has chosen a
QUALITATIVE research approach. Given the research question, theory, and construct map,
design a rigorous qualitative methodology and generate a semi-structured interview protocol.

Your output must include:
1. Recommended qualitative method (phenomenological interviews, grounded theory,
   case study, ethnography, narrative inquiry) with rationale
2. Sampling strategy with target sample size and inclusion/exclusion criteria
3. A semi-structured interview protocol with:
   - 3-5 thematic sections aligned with the research question and constructs
   - 4-6 open-ended questions per section
   - 2-3 probing/follow-up questions for each main question
   - An opening ice-breaker and closing reflection question
4. Data analysis approach and trustworthiness criteria

Questions must be open-ended, non-leading, theoretically grounded.
Use "tell me about..." and "how do you..." stems. No yes/no questions.

Return ONLY valid JSON, no fences:
{
  "recommended_method": "...",
  "method_rationale": "...",
  "sample_plan": {
    "strategy": "...", "target_n": 0, "saturation_plan": "...",
    "recruitment": "...", "inclusion_criteria": "..."
  },
  "interview_protocol": {
    "sections": [
      {
        "section_name": "...", "construct_explored": "...",
        "questions": [{"id": "Q1", "text": "...", "probes": ["...", "..."]}]
      }
    ],
    "opening_question": "...", "closing_question": "...",
    "estimated_duration_minutes": 45
  },
  "analysis_approach": {
    "method": "...", "coding_strategy": "...",
    "trustworthiness": ["...", "..."]
  },
  "ethics_flags": ["..."]
}
"""

SYS_FULL_PAPER = """You are the Paper Architect agent in Eagle. You have the ENTIRE research
canvas including the reviewer critique. Auto-fix ALL reviewer concerns and
generate a publication-ready paper with fully written sections.

Write in academic prose for a top management/marketing journal (APA style).
Do NOT hallucinate references — use only those from the lit table. Where thin,
note "[additional references needed]."

Generate these sections:
1. INTRODUCTION (500-700 words): Hook, tension/gap, theoretical contribution,
   RQ, paper structure preview.
2. LITERATURE REVIEW (600-900 words): Organize by themes, define constructs,
   build toward gap. End with "the present study addresses..."
3. HYPOTHESES: Each hypothesis with theoretical justification (100-150 words each).
4. METHOD (400-600 words): Design, sample, procedure, measurement. If qualitative,
   describe interview protocol and analysis approach.
5. THEORETICAL IMPLICATIONS: Exactly 4, each 80-120 words.
6. PRACTICAL IMPLICATIONS: Exactly 4, each 80-120 words with actionable recommendations.

Address ALL reviewer concerns in appropriate sections.

Return ONLY valid JSON, no fences:
{
  "introduction": "...",
  "literature_review": "...",
  "hypotheses_text": "...",
  "method_section": "...",
  "theoretical_implications": [
    {"id": "TI1", "title": "...", "text": "..."},
    {"id": "TI2", "title": "...", "text": "..."},
    {"id": "TI3", "title": "...", "text": "..."},
    {"id": "TI4", "title": "...", "text": "..."}
  ],
  "practical_implications": [
    {"id": "PI1", "title": "...", "text": "..."},
    {"id": "PI2", "title": "...", "text": "..."},
    {"id": "PI3", "title": "...", "text": "..."},
    {"id": "PI4", "title": "...", "text": "..."}
  ],
  "revised_notes": "Summary of how reviewer concerns were addressed..."
}
"""

# ---------------------------------------------------------------------------
# WORD EXPORT
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
GOLD = RGBColor(0xC8, 0xA2, 0x4C)

def export_docx(canvas: dict) -> bytes:
    doc = Document()
    # title
    t = doc.add_heading("Research Canvas", level=0)
    for r in t.runs:
        r.font.color.rgb = NAVY

    meta = doc.add_paragraph()
    meta.add_run(f"Faculty: {canvas['faculty_id']}    ").italic = True
    meta.add_run(f"Session: {canvas['session_id'][:8]}    ").italic = True
    meta.add_run(f"Updated: {canvas['updated_at'][:19]}").italic = True

    # Stage 1
    s1 = canvas["stage_1_problematization"]
    doc.add_heading("1. Problematization", level=1)
    doc.add_paragraph(f"Fuzzy idea: {s1.get('fuzzy_idea','')}")
    doc.add_paragraph(f"Selected RQ: {s1.get('selected_rq','')}")

    # Stage 2
    s2 = canvas["stage_2_theorization"]
    doc.add_heading("2. Theorization", level=1)
    doc.add_paragraph(f"Theory: {s2.get('selected_theory','')}")
    cm = s2.get("construct_map", {}) or {}
    if cm:
        doc.add_paragraph("Construct Map:", style="Intense Quote")
        for key in ["focal_iv", "focal_dv", "mediators", "moderators"]:
            items = cm.get(key, []) or []
            if items:
                doc.add_paragraph(f"{key.replace('_',' ').title()}: " +
                                  "; ".join(i.get("name","") for i in items))
        if cm.get("controls"):
            doc.add_paragraph("Controls: " + ", ".join(cm["controls"]))
        if cm.get("boundary_conditions"):
            doc.add_paragraph(f"Boundary: {cm['boundary_conditions']}")
    for h in s2.get("hypotheses", []) or []:
        doc.add_paragraph(f"{h.get('id','H')}: {h.get('text','')}", style="List Bullet")

    # Stage 3 — lit table
    s3 = canvas["stage_3_literature"]
    doc.add_heading("3. Literature Review", level=1)
    lit = s3.get("lit_table", []) or []
    if lit:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Authors (Year)", "Venue", "Theory/Method", "Findings", "Gap"]):
            hdr[i].text = h
        for p in lit:
            row = table.add_row().cells
            row[0].text = f"{p.get('authors','')} ({p.get('year','')})"
            row[1].text = p.get("venue","")
            row[2].text = f"{p.get('theory','')} / {p.get('method','')}"
            row[3].text = p.get("findings","")
            row[4].text = p.get("gap_identified","")
    if s3.get("synthesis"):
        doc.add_heading("Synthesis", level=2)
        doc.add_paragraph(s3["synthesis"])
    if s3.get("identified_gap"):
        doc.add_paragraph(f"Gap: {s3['identified_gap']}", style="Intense Quote")

    # Stage 4 — method & survey
    s4 = canvas["stage_4_method"]
    doc.add_heading("4. Method & Survey Instrument", level=1)
    doc.add_paragraph(f"Method: {s4.get('recommended_method','')}")
    doc.add_paragraph(f"Rationale: {s4.get('method_rationale','')}")
    sp = s4.get("sample_plan", {}) or {}
    if sp:
        doc.add_paragraph(f"Target N: {sp.get('target_n','')}   |   Recruitment: {sp.get('recruitment','')}")
        if sp.get("power_analysis"):
            doc.add_paragraph(f"Power: {sp['power_analysis']}")
    instr = s4.get("instrument", {}) or {}
    for block in instr.get("blocks", []) or []:
        doc.add_heading(block.get("block_name", "Block"), level=2)
        doc.add_paragraph(f"Construct: {block.get('construct','')}   |   "
                          f"Source: {block.get('scale_source','')}   |   "
                          f"Validated: {block.get('validated','')}")
        for item in block.get("items", []) or []:
            tag = " (R)" if item.get("reverse_coded") else ""
            doc.add_paragraph(f"{item.get('id','')}{tag}: {item.get('text','')}",
                              style="List Number")
    if instr.get("attention_checks"):
        doc.add_heading("Attention Checks", level=2)
        for ac in instr["attention_checks"]:
            doc.add_paragraph(f"{ac.get('id','')}: {ac.get('text','')}", style="List Bullet")
    if instr.get("demographics"):
        doc.add_paragraph("Demographics: " + ", ".join(instr["demographics"]))

    # Stage 5 — review
    s5 = canvas["stage_5_review"]
    doc.add_heading("5. Reviewer Critique", level=1)
    doc.add_paragraph(f"Target journal: {s5.get('target_journal','')}   |   "
                      f"Readiness: {s5.get('overall_readiness_score','')}/10")
    for c in s5.get("reviewer_critique", []) or []:
        p = doc.add_paragraph()
        p.add_run(f"[{c.get('severity','').upper()}] {c.get('category','')} — ").bold = True
        p.add_run(c.get("issue",""))
        doc.add_paragraph(f"Fix: {c.get('suggested_fix','')}", style="Intense Quote")
    if s5.get("journal_specific_notes"):
        doc.add_paragraph(s5["journal_specific_notes"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ---------------------------------------------------------------------------
# STREAMLIT UI
# ---------------------------------------------------------------------------

st.set_page_config(page_title="Eagle — Research Copilot", layout="wide",
                   page_icon="🦅")

# ---------------------------------------------------------------------------
# PREMIUM LIGHT THEME CSS
# ---------------------------------------------------------------------------
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');

/* ── Global ── */
html, body, [class*="css"] {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
}
.stApp {
    background: linear-gradient(160deg, #f8f9fc 0%, #eef1f8 50%, #f0f0fa 100%);
    color: #1e293b;
}

/* ── Sidebar ── */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #ffffff 0%, #f8f9fc 100%);
    border-right: 1px solid #e2e8f0;
    box-shadow: 4px 0 24px rgba(0,0,0,0.03);
}
section[data-testid="stSidebar"] .stMarkdown h2 {
    background: linear-gradient(135deg, #4338ca, #6366f1, #818cf8);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    font-weight: 900;
    font-size: 1.7rem;
    letter-spacing: -0.03em;
}
section[data-testid="stSidebar"] .stCaption {
    color: #94a3b8;
    font-size: 0.7rem;
    text-transform: uppercase;
    letter-spacing: 0.1em;
    font-weight: 600;
}

/* ── Inputs ── */
.stTextInput input, .stTextArea textarea, .stSelectbox > div > div {
    background: #ffffff !important;
    border: 1.5px solid #e2e8f0 !important;
    border-radius: 12px !important;
    color: #1e293b !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 0.9rem !important;
    transition: all 0.25s cubic-bezier(0.4,0,0.2,1);
    box-shadow: 0 1px 3px rgba(0,0,0,0.04) !important;
}
.stTextInput input:focus, .stTextArea textarea:focus {
    border-color: #818cf8 !important;
    box-shadow: 0 0 0 3px rgba(99,102,241,0.12), 0 1px 3px rgba(0,0,0,0.04) !important;
}

/* ── Primary Buttons ── */
.stButton > button {
    background: linear-gradient(135deg, #4f46e5 0%, #6366f1 100%) !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 12px !important;
    font-weight: 600 !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 0.875rem !important;
    padding: 0.6rem 1.5rem !important;
    transition: all 0.3s cubic-bezier(0.4,0,0.2,1) !important;
    box-shadow: 0 2px 8px rgba(79,70,229,0.25), 0 1px 2px rgba(0,0,0,0.05) !important;
    letter-spacing: 0.01em;
}
.stButton > button:hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 6px 20px rgba(79,70,229,0.3), 0 2px 4px rgba(0,0,0,0.05) !important;
    background: linear-gradient(135deg, #4338ca 0%, #4f46e5 100%) !important;
}
.stButton > button:active {
    transform: translateY(0) !important;
}
.stDownloadButton > button {
    background: linear-gradient(135deg, #059669 0%, #10b981 100%) !important;
    box-shadow: 0 2px 8px rgba(5,150,105,0.25) !important;
}
.stDownloadButton > button:hover {
    box-shadow: 0 6px 20px rgba(5,150,105,0.3) !important;
    transform: translateY(-1px) !important;
}

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] {
    gap: 6px;
    background: #ffffff;
    border-radius: 16px;
    padding: 6px;
    border: 1px solid #e2e8f0;
    box-shadow: 0 1px 4px rgba(0,0,0,0.04);
}
.stTabs [data-baseweb="tab"] {
    border-radius: 12px;
    padding: 10px 22px;
    font-weight: 500;
    color: #64748b;
    font-family: 'Inter', sans-serif;
    font-size: 0.85rem;
    transition: all 0.25s ease;
    border-bottom: none !important;
}
.stTabs [data-baseweb="tab"]:hover {
    color: #4f46e5;
    background: #f1f5f9;
}
.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #4f46e5, #6366f1) !important;
    color: #ffffff !important;
    font-weight: 600;
    border-bottom-color: transparent !important;
    box-shadow: 0 2px 8px rgba(79,70,229,0.25);
}

/* ── Expanders ── */
.streamlit-expanderHeader {
    background: #ffffff !important;
    border: 1px solid #e2e8f0 !important;
    border-radius: 12px !important;
    font-weight: 500 !important;
    color: #1e293b !important;
    transition: all 0.25s ease;
    box-shadow: 0 1px 3px rgba(0,0,0,0.03);
}
.streamlit-expanderHeader:hover {
    border-color: #c7d2fe !important;
    background: #fafafe !important;
    box-shadow: 0 2px 8px rgba(99,102,241,0.08);
}
details[open] .streamlit-expanderHeader {
    border-color: #818cf8 !important;
}

/* ── Metrics ── */
[data-testid="stMetric"] {
    background: linear-gradient(135deg, #eef2ff, #e0e7ff);
    border: 1px solid #c7d2fe;
    border-radius: 16px;
    padding: 1.2rem 1.5rem;
    box-shadow: 0 2px 8px rgba(99,102,241,0.06);
}
[data-testid="stMetricValue"] {
    color: #4338ca !important;
    font-weight: 800;
    font-size: 2rem !important;
}
[data-testid="stMetricLabel"] {
    color: #6366f1 !important;
    font-weight: 600;
    text-transform: uppercase;
    font-size: 0.7rem !important;
    letter-spacing: 0.05em;
}

/* ── Headings ── */
h1 {
    background: linear-gradient(135deg, #312e81, #4338ca, #4f46e5);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    font-weight: 900 !important;
    letter-spacing: -0.04em;
    line-height: 1.1 !important;
}
h2 {
    color: #312e81 !important;
    font-weight: 800 !important;
    letter-spacing: -0.02em;
    font-size: 1.4rem !important;
}
h3 {
    color: #4338ca !important;
    font-weight: 700 !important;
    letter-spacing: -0.01em;
}
.stSubheader {
    color: #1e293b !important;
}

/* ── Alerts ── */
.stAlert {
    border-radius: 12px !important;
    border-left-width: 4px !important;
    font-size: 0.88rem;
}

/* ── Dataframe ── */
.stDataFrame {
    border-radius: 12px;
    overflow: hidden;
    border: 1px solid #e2e8f0;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
}

/* ── Dividers ── */
hr {
    border-color: #e2e8f0 !important;
    opacity: 0.6;
}

/* ── Blockquote (Socratic answers) ── */
blockquote {
    border-left: 3px solid #6366f1 !important;
    background: linear-gradient(135deg, #f5f3ff, #eef2ff);
    padding: 0.75rem 1rem;
    border-radius: 0 12px 12px 0;
    color: #374151;
    font-style: normal !important;
    margin: 0.5rem 0;
}

/* ── JSON viewer ── */
.stJson {
    background: #ffffff !important;
    border: 1px solid #e2e8f0;
    border-radius: 12px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.04);
}

/* ── Radio ── */
.stRadio > label {
    color: #374151 !important;
    font-weight: 500;
}

/* ── Progress bar ── */
.stProgress > div > div {
    background: linear-gradient(90deg, #4f46e5, #818cf8) !important;
    border-radius: 8px;
}

/* ── Caption ── */
.stCaption {
    color: #64748b !important;
}

/* ── Strong text ── */
strong {
    color: #1e293b;
}

/* ── Selection ── */
::selection {
    background: rgba(99,102,241,0.15);
    color: #312e81;
}

/* ── Spinner ── */
.stSpinner > div > div {
    border-top-color: #6366f1 !important;
}

/* ── Success/Info message enhancements ── */
.stSuccess {
    background: linear-gradient(135deg, #f0fdf4, #ecfdf5) !important;
    border-color: #86efac !important;
}
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# DB MIGRATION — carry over old sherpa.db data if present
# ---------------------------------------------------------------------------
import os
if os.path.exists("sherpa.db") and not os.path.exists(DB_PATH):
    import shutil
    shutil.copy2("sherpa.db", DB_PATH)

# session state
db_init()
if "canvas" not in st.session_state:
    st.session_state.canvas = None
if "faculty_id" not in st.session_state:
    st.session_state.faculty_id = ""

# ---------- SIDEBAR ----------
with st.sidebar:
    st.markdown('<h2 style="margin:0;">🦅 Eagle</h2>', unsafe_allow_html=True)
    st.caption("Research Copilot for GLIM Faculty")
    st.divider()

    st.session_state.anthropic_api_key = st.text_input(
        "Anthropic API Key", type="password",
        value=st.session_state.get("anthropic_api_key", ""),
        placeholder="sk-ant-..."
    )

    st.session_state.faculty_id = st.text_input(
        "Faculty email", value=st.session_state.faculty_id or "harish@glim.ac.in"
    )

    if st.button("➕ New session", use_container_width=True):
        st.session_state.canvas = new_canvas(st.session_state.faculty_id)
        db_save(st.session_state.canvas)
        st.rerun()

    st.divider()
    st.caption("Past sessions")
    if st.session_state.faculty_id:
        for sid, title, updated in db_list(st.session_state.faculty_id):
            if st.button(f"📄 {title[:30]}", key=f"load_{sid}", use_container_width=True):
                st.session_state.canvas = db_load(sid)
                st.rerun()

    st.divider()
    if st.session_state.canvas:
        docx_bytes = export_docx(st.session_state.canvas)
        st.download_button(
            "⬇️ Export to Word",
            data=docx_bytes,
            file_name=f"eagle_{st.session_state.canvas['session_id'][:8]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

# ---------- MAIN ----------
if not st.session_state.canvas:
    st.title("🧪 Eagle — Research Copilot")
    st.markdown("""
    **A research partner for faculty. Not a chatbot.**

    Eagle walks you through the five hardest parts of kicking off a serious
    research project:
    1. **Problematize** — sharpen a fuzzy idea into a research question
    2. **Theorize** — pick a lens and build a construct map
    3. **Literature** — real papers from OpenAlex, structured into a gap-finding table
    4. **Method & Survey** — recommended design and a publication-ready instrument
    5. **Reviewer** — devil's advocate critique at the target journal level

    Start a new session in the sidebar.
    """)
    st.stop()

canvas = st.session_state.canvas

# Title editor
new_title = st.text_input("Session title", value=canvas.get("title", ""),
                          label_visibility="collapsed")
if new_title != canvas.get("title"):
    canvas["title"] = new_title
    db_save(canvas)

# Stage tabs
tabs = st.tabs([label for _, label in STAGES])

# ===== STAGE 1 — PROBLEMATIZER =====
with tabs[0]:
    s1 = canvas["stage_1_problematization"]
    st.subheader("Start with the fuzzy idea")
    st.caption("One or two sentences. Don't polish.")

    fuzzy = st.text_area("Fuzzy idea", value=s1["fuzzy_idea"], height=80,
                         label_visibility="collapsed")
    if fuzzy != s1["fuzzy_idea"]:
        s1["fuzzy_idea"] = fuzzy
        db_save(canvas)

    st.divider()
    st.markdown("### Socratic exchange")
    for i, ex in enumerate(s1["socratic_exchange"]):
        st.markdown(f"**Q{i+1}.** {ex['q']}")
        st.markdown(f"> {ex['a']}")

    # next question
    if s1["fuzzy_idea"] and len(s1["socratic_exchange"]) < 6 and not s1["candidate_rqs"]:
        col1, col2 = st.columns([3, 1])
        with col1:
            if st.button("Ask next Socratic question", type="primary"):
                history = f"Fuzzy idea: {s1['fuzzy_idea']}\n\nExchange so far:\n"
                for ex in s1["socratic_exchange"]:
                    history += f"Q: {ex['q']}\nA: {ex['a']}\n"
                with st.spinner("Thinking..."):
                    q = claude_call(SYS_PROBLEMATIZER_QUESTION, history)
                s1["socratic_exchange"].append({"q": q.strip(), "a": ""})
                db_save(canvas)
                st.rerun()

    # answer any unanswered question (find the first one)
    unanswered_idx = None
    for idx, ex in enumerate(s1["socratic_exchange"]):
        if not ex["a"]:
            unanswered_idx = idx
            break
    if unanswered_idx is not None:
        st.markdown(f"**Q{unanswered_idx+1}.** {s1['socratic_exchange'][unanswered_idx]['q']}")
        ans = st.text_area("Your answer",
                           key=f"ans_{unanswered_idx}", height=100)
        if st.button("Submit answer"):
            s1["socratic_exchange"][unanswered_idx]["a"] = ans
            db_save(canvas)
            st.rerun()

    # finalize RQs
    if (len(s1["socratic_exchange"]) >= 3
        and all(ex["a"] for ex in s1["socratic_exchange"])
        and not s1["candidate_rqs"]):
        if st.button("📝 Generate candidate research questions", type="primary"):
            history = f"Fuzzy idea: {s1['fuzzy_idea']}\n\nExchange:\n"
            for ex in s1["socratic_exchange"]:
                history += f"Q: {ex['q']}\nA: {ex['a']}\n"
            with st.spinner("Drafting candidate RQs..."):
                out = claude_json(SYS_PROBLEMATIZER_FINALIZE, history)
            s1["candidate_rqs"] = out["candidate_rqs"]
            db_save(canvas)
            st.rerun()

    if s1["candidate_rqs"]:
        st.markdown("### Candidate research questions")
        for i, cand in enumerate(s1["candidate_rqs"]):
            st.markdown(f"**Option {i+1}.** {cand['rq']}")
            st.caption(f"Trade-off: {cand['tradeoff']}")
        choice = st.radio("Select one", options=[c["rq"] for c in s1["candidate_rqs"]],
                          index=None)
        if choice and choice != s1["selected_rq"]:
            s1["selected_rq"] = choice
            db_save(canvas)
            st.success("RQ locked. Move to Stage 2 — Theorize.")

# ===== STAGE 2 — THEORIST =====
with tabs[1]:
    s1 = canvas["stage_1_problematization"]
    s2 = canvas["stage_2_theorization"]

    if not s1["selected_rq"]:
        st.info("Complete Stage 1 first.")
    else:
        st.markdown(f"**RQ:** {s1['selected_rq']}")
        if not s2["candidate_theories"]:
            if st.button("🧠 Propose theories & construct map", type="primary"):
                with st.spinner("Theorist working..."):
                    out = claude_json(SYS_THEORIST, f"RQ: {s1['selected_rq']}")
                s2["candidate_theories"] = out.get("candidate_theories", [])
                s2["selected_theory"] = out.get("recommended_theory", "")
                s2["construct_map"] = out.get("construct_map", {})
                s2["hypotheses"] = out.get("hypotheses", [])
                db_save(canvas)
                st.rerun()
        else:
            st.markdown("### Candidate theories")
            for t in s2["candidate_theories"]:
                with st.expander(f"📚 {t['name']}"):
                    st.write(f"**Core claim:** {t.get('core_claim','')}")
                    st.write(f"**Strengths:** {t.get('strengths','')}")
                    st.write(f"**Weaknesses:** {t.get('weaknesses','')}")
                    refs = t.get("seminal_refs", [])
                    if refs:
                        st.write("**Seminal:** " + "; ".join(refs))

            s2["selected_theory"] = st.text_input("Selected theory (edit if needed)",
                                                  value=s2["selected_theory"])
            st.markdown("### Construct map")
            st.json(s2["construct_map"])
            st.markdown("### Hypotheses")
            for h in s2["hypotheses"]:
                st.markdown(f"- **{h['id']}** ({h.get('type','')}): {h['text']}")
            if st.button("💾 Save & continue to Stage 3"):
                db_save(canvas)
                st.success("Saved.")

# ===== STAGE 3 — LITERATURE =====
with tabs[2]:
    s2 = canvas["stage_2_theorization"]
    s3 = canvas["stage_3_literature"]

    if not s2.get("construct_map"):
        st.info("Complete Stage 2 first.")
    else:
        if not s3["lit_table"]:
            if st.button("🔍 Search OpenAlex & build lit review", type="primary"):
                # Step 1: generate queries
                context = json.dumps({
                    "theory": s2["selected_theory"],
                    "construct_map": s2["construct_map"],
                    "rq": canvas["stage_1_problematization"]["selected_rq"],
                })
                with st.spinner("Generating search queries..."):
                    q_out = claude_json(SYS_LITSCOUT_QUERIES, context, model=MODEL_LIGHT)
                queries = q_out.get("queries", [])
                s3["search_queries_used"] = queries

                # Step 2: hit OpenAlex
                all_papers = []
                seen = set()
                progress = st.progress(0.0, "Querying OpenAlex...")
                for i, q in enumerate(queries):
                    papers = openalex_search(q, per_page=20)
                    for p in papers:
                        if p["openalex_id"] not in seen:
                            seen.add(p["openalex_id"])
                            all_papers.append(p)
                    progress.progress((i + 1) / len(queries), f"Querying: {q}")
                progress.empty()
                st.info(f"Retrieved {len(all_papers)} unique papers. Synthesizing...")

                # Step 3: send to Claude for scoring + synthesis
                # keep top 40 by citations to stay within context
                all_papers.sort(key=lambda p: p.get("cited_by", 0), reverse=True)
                paper_blob = json.dumps(all_papers[:25])
                synth_input = (
                    f"RQ: {canvas['stage_1_problematization']['selected_rq']}\n"
                    f"Theory: {s2['selected_theory']}\n"
                    f"Construct map: {json.dumps(s2['construct_map'])}\n\n"
                    f"Papers from OpenAlex:\n{paper_blob}"
                )
                with st.spinner("Claude building lit table and synthesis..."):
                    out = claude_json(SYS_LITSCOUT_SYNTHESIZE, synth_input)
                s3["lit_table"] = out.get("lit_table", [])
                s3["synthesis"] = out.get("synthesis", "")
                s3["identified_gap"] = out.get("identified_gap", "")
                db_save(canvas)
                st.rerun()
        else:
            st.markdown(f"**Queries used:** {', '.join(s3['search_queries_used'])}")
            st.markdown("### Literature table")
            st.dataframe(s3["lit_table"], use_container_width=True)
            st.markdown("### Synthesis")
            st.write(s3["synthesis"])
            st.markdown("### Identified gap")
            st.info(s3["identified_gap"])
            if st.button("🔄 Re-run literature search"):
                s3["lit_table"] = []
                db_save(canvas)
                st.rerun()

# ===== STAGE 4 — METHOD =====
with tabs[3]:
    s3 = canvas["stage_3_literature"]
    s4 = canvas["stage_4_method"]

    if not s3.get("lit_table"):
        st.info("Complete Stage 3 first.")
    else:
        # Research approach toggle
        approach_options = ["quantitative", "qualitative"]
        current_approach = s4.get("research_approach", "quantitative")
        selected_approach = st.radio(
            "Research approach",
            approach_options,
            index=approach_options.index(current_approach) if current_approach in approach_options else 0,
            horizontal=True,
            key="research_approach_radio"
        )
        if selected_approach != s4.get("research_approach"):
            s4["research_approach"] = selected_approach
            db_save(canvas)

        st.divider()
        is_qual = s4["research_approach"] == "qualitative"
        method_ready = (s4["interview_protocol"].get("sections") if is_qual
                        else s4["instrument"]["blocks"])

        if not method_ready:
            btn_label = "🎙️ Design qualitative method & interview protocol" if is_qual else "🧪 Design method & generate survey"
            if st.button(btn_label, type="primary"):
                context = json.dumps({
                    "rq": canvas["stage_1_problematization"]["selected_rq"],
                    "theory": canvas["stage_2_theorization"]["selected_theory"],
                    "construct_map": canvas["stage_2_theorization"]["construct_map"],
                    "hypotheses": canvas["stage_2_theorization"]["hypotheses"],
                    "gap": s3["identified_gap"],
                })
                if is_qual:
                    with st.spinner("Designing qualitative method & interview protocol..."):
                        out = claude_json(SYS_METHOD_QUALITATIVE, context)
                    s4["recommended_method"] = out.get("recommended_method", "")
                    s4["method_rationale"] = out.get("method_rationale", "")
                    s4["sample_plan"] = out.get("sample_plan", {})
                    s4["interview_protocol"] = out.get("interview_protocol", {})
                    s4["ethics_flags"] = out.get("ethics_flags", [])
                    if out.get("analysis_approach"):
                        s4["analysis_approach"] = out["analysis_approach"]
                else:
                    with st.spinner("Designing method and instrument..."):
                        out = claude_json(SYS_METHOD, context)
                    s4["recommended_method"] = out.get("recommended_method", "")
                    s4["method_rationale"] = out.get("method_rationale", "")
                    s4["sample_plan"] = out.get("sample_plan", {})
                    s4["instrument"] = out.get("instrument", {})
                    s4["ethics_flags"] = out.get("ethics_flags", [])
                db_save(canvas)
                st.rerun()
        else:
            st.markdown(f"**Method:** {s4['recommended_method']}")
            st.caption(s4["method_rationale"])
            sp = s4["sample_plan"]

            if is_qual:
                # ── Qualitative display ──
                st.markdown(f"**Sampling:** {sp.get('strategy','')}   |   N={sp.get('target_n','?')}   |   {sp.get('recruitment','')}")
                if sp.get("saturation_plan"):
                    st.caption(f"Saturation: {sp['saturation_plan']}")

                st.markdown("### 🎙️ Interview Protocol")
                ip = s4["interview_protocol"]
                if ip.get("opening_question"):
                    st.info(f"**Opening:** {ip['opening_question']}")
                for sec in ip.get("sections", []):
                    with st.expander(f"📝 {sec['section_name']} — {sec.get('construct_explored','')}"):
                        for q in sec.get("questions", []):
                            st.markdown(f"**{q['id']}.** {q['text']}")
                            if q.get("probes"):
                                for probe in q["probes"]:
                                    st.caption(f"   ↳ Probe: {probe}")
                if ip.get("closing_question"):
                    st.info(f"**Closing:** {ip['closing_question']}")
                if ip.get("estimated_duration_minutes"):
                    st.caption(f"Estimated duration: {ip['estimated_duration_minutes']} minutes")

                # Analysis approach
                aa = s4.get("analysis_approach", {})
                if aa:
                    st.markdown("### Analysis Approach")
                    st.markdown(f"**Method:** {aa.get('method','')}")
                    st.markdown(f"**Coding:** {aa.get('coding_strategy','')}")
                    if aa.get("trustworthiness"):
                        st.markdown("**Trustworthiness:** " + ", ".join(aa["trustworthiness"]))
            else:
                # ── Quantitative display ──
                st.markdown(f"**Sample:** N={sp.get('target_n','?')}   |   {sp.get('recruitment','')}")
                st.caption(sp.get("power_analysis", ""))

                st.markdown("### Survey instrument")
                for block in s4["instrument"]["blocks"]:
                    with st.expander(f"📋 {block['block_name']} — {block.get('construct','')}"):
                        st.caption(f"Source: {block.get('scale_source','')}   |   "
                                   f"Validated: {block.get('validated','')}")
                        for item in block.get("items", []):
                            tag = " (R)" if item.get("reverse_coded") else ""
                            st.markdown(f"- **{item['id']}**{tag}: {item['text']}")

                if s4["instrument"].get("attention_checks"):
                    st.markdown("**Attention checks:** " +
                                "; ".join(a["text"] for a in s4["instrument"]["attention_checks"]))

            if s4["ethics_flags"]:
                st.warning("Ethics flags: " + "; ".join(s4["ethics_flags"]))
            if st.button("🔄 Re-run method design"):
                s4["instrument"] = {"blocks": [], "attention_checks": [], "demographics": []}
                s4["interview_protocol"] = {"sections": [], "probes": [], "sampling_strategy": ""}
                db_save(canvas)
                st.rerun()

# ===== STAGE 5 — REVIEWER =====
with tabs[4]:
    s5 = canvas["stage_5_review"]
    s4 = canvas["stage_4_method"]
    is_qual = s4.get("research_approach") == "qualitative"
    method_done = (s4["interview_protocol"].get("sections") if is_qual
                   else s4["instrument"]["blocks"])
    if not method_done:
        st.info("Complete Stage 4 first.")
    else:
        s5["target_journal"] = st.selectbox(
            "Target journal",
            ["JAMS", "JCR", "JMR", "MISQ", "MIT Sloan Mgmt Review", "Information & Management", "IJRDM"],
            index=0 if not s5["target_journal"] else
                  ["JAMS","JCR","JMR","MISQ","MIT Sloan Mgmt Review","Information & Management","IJRDM"].index(s5["target_journal"])
                  if s5["target_journal"] in ["JAMS","JCR","JMR","MISQ","MIT Sloan Mgmt Review","Information & Management","IJRDM"] else 0
        )

        if not s5["reviewer_critique"]:
            if st.button("🎯 Run Reviewer critique", type="primary"):
                # send the whole canvas minus stage_5 itself
                review_canvas = {k: v for k, v in canvas.items() if k != "stage_5_review"}
                review_canvas["target_journal"] = s5["target_journal"]
                with st.spinner(f"Reviewing for {s5['target_journal']}..."):
                    out = claude_json(SYS_REVIEWER, json.dumps(review_canvas))
                s5["reviewer_critique"] = out.get("reviewer_critique", [])
                s5["overall_readiness_score"] = out.get("overall_readiness_score", 0)
                s5["journal_specific_notes"] = out.get("journal_specific_notes", "")
                db_save(canvas)
                st.rerun()
        else:
            score = s5["overall_readiness_score"]
            col1, col2 = st.columns([1, 3])
            with col1:
                st.metric(f"Readiness for {s5['target_journal']}", f"{score}/10")
            with col2:
                st.caption(s5["journal_specific_notes"])

            for c in s5["reviewer_critique"]:
                sev = c.get("severity", "medium")
                icon = {"high": "🔴", "medium": "🟡", "low": "🟢"}.get(sev, "⚪")
                with st.expander(f"{icon} {c['concern_id']} — {c['category']}"):
                    st.markdown(f"**Issue:** {c['issue']}")
                    st.markdown(f"**Fix:** {c['suggested_fix']}")
                    st.caption(f"Revise in: {c.get('linked_stage','')}")
            if st.button("🔄 Re-run review"):
                s5["reviewer_critique"] = []
                db_save(canvas)
                st.rerun()

# ===== STAGE 6 — FULL PAPER =====
with tabs[5]:
    s5 = canvas["stage_5_review"]
    s6 = canvas.get("stage_6_full_paper", {})
    # Ensure stage_6 exists in canvas for older sessions
    if "stage_6_full_paper" not in canvas:
        canvas["stage_6_full_paper"] = {
            "introduction": "", "literature_review": "", "hypotheses_text": "",
            "method_section": "", "theoretical_implications": [],
            "practical_implications": [], "revised_notes": "", "generated": False,
        }
        s6 = canvas["stage_6_full_paper"]

    if not s5.get("reviewer_critique"):
        st.info("Complete Stage 5 (Reviewer Critique) first. The full paper generator uses reviewer feedback to auto-fix and improve everything.")
    else:
        if not s6.get("generated"):
            st.markdown("### 📄 Generate Full Paper")
            st.markdown("""
            Based on the reviewer's critique, Eagle will **auto-fix all concerns** and generate
            a publication-ready paper with:
            - **Introduction** (500-700 words)
            - **Literature Review** (600-900 words)
            - **Hypotheses** with theoretical justification
            - **Method** section
            - **4 Theoretical Implications**
            - **4 Practical Implications**
            """)
            if st.button("🚀 Generate Full Paper (auto-fix & write)", type="primary"):
                # Send the entire canvas to the Paper Architect
                paper_context = json.dumps({
                    "rq": canvas["stage_1_problematization"]["selected_rq"],
                    "theory": canvas["stage_2_theorization"]["selected_theory"],
                    "construct_map": canvas["stage_2_theorization"]["construct_map"],
                    "hypotheses": canvas["stage_2_theorization"]["hypotheses"],
                    "lit_table": canvas["stage_3_literature"]["lit_table"],
                    "synthesis": canvas["stage_3_literature"]["synthesis"],
                    "identified_gap": canvas["stage_3_literature"]["identified_gap"],
                    "research_approach": canvas["stage_4_method"].get("research_approach", "quantitative"),
                    "method": canvas["stage_4_method"]["recommended_method"],
                    "method_rationale": canvas["stage_4_method"]["method_rationale"],
                    "sample_plan": canvas["stage_4_method"]["sample_plan"],
                    "reviewer_critique": s5["reviewer_critique"],
                    "overall_readiness_score": s5["overall_readiness_score"],
                    "journal_specific_notes": s5.get("journal_specific_notes", ""),
                    "target_journal": s5.get("target_journal", "JAMS"),
                })
                with st.spinner("🧠 Eagle is auto-fixing reviewer concerns and writing your full paper... This may take a minute."):
                    out = claude_json(SYS_FULL_PAPER, paper_context, max_tokens=8192)
                s6["introduction"] = out.get("introduction", "")
                s6["literature_review"] = out.get("literature_review", "")
                s6["hypotheses_text"] = out.get("hypotheses_text", "")
                s6["method_section"] = out.get("method_section", "")
                s6["theoretical_implications"] = out.get("theoretical_implications", [])
                s6["practical_implications"] = out.get("practical_implications", [])
                s6["revised_notes"] = out.get("revised_notes", "")
                s6["generated"] = True
                db_save(canvas)
                st.rerun()
        else:
            st.success("✅ Full paper generated with all reviewer fixes applied!")
            if s6.get("revised_notes"):
                with st.expander("📋 Reviewer fixes applied"):
                    st.write(s6["revised_notes"])

            st.markdown("### 1. Introduction")
            st.write(s6["introduction"])

            st.markdown("### 2. Literature Review")
            st.write(s6["literature_review"])

            st.markdown("### 3. Hypotheses")
            st.write(s6["hypotheses_text"])

            st.markdown("### 4. Method")
            st.write(s6["method_section"])

            st.markdown("### 5. Theoretical Implications")
            for imp in s6.get("theoretical_implications", []):
                with st.expander(f"🔬 {imp.get('id','')} — {imp.get('title','')}"):
                    st.write(imp.get("text", ""))

            st.markdown("### 6. Practical Implications")
            for imp in s6.get("practical_implications", []):
                with st.expander(f"💼 {imp.get('id','')} — {imp.get('title','')}"):
                    st.write(imp.get("text", ""))

            # Full paper Word export
            st.divider()
            if st.button("📥 Generate Full Paper Word Document", type="primary"):
                docx_bytes = export_full_paper_docx(canvas)
                st.download_button(
                    "⬇️ Download Full Paper (.docx)",
                    data=docx_bytes,
                    file_name=f"eagle_paper_{canvas['session_id'][:8]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

            if st.button("🔄 Re-generate full paper"):
                s6["generated"] = False
                s6["introduction"] = ""
                db_save(canvas)
                st.rerun()


def export_full_paper_docx(canvas: dict) -> bytes:
    """Generate a professionally formatted Word document with the full paper."""
    doc = Document()
    INDIGO = RGBColor(0x4F, 0x46, 0xE5)
    DARK = RGBColor(0x1E, 0x29, 0x3B)
    s6 = canvas.get("stage_6_full_paper", {})

    # Title
    title = doc.add_paragraph()
    run = title.add_run(canvas.get("title", "Research Paper"))
    run.font.size = Pt(24)
    run.font.color.rgb = INDIGO
    run.bold = True

    # RQ
    rq = canvas["stage_1_problematization"]["selected_rq"]
    p = doc.add_paragraph()
    run = p.add_run(f"Research Question: {rq}")
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = DARK

    doc.add_paragraph()  # spacer

    # Sections
    sections_data = [
        ("1. INTRODUCTION", s6.get("introduction", "")),
        ("2. LITERATURE REVIEW", s6.get("literature_review", "")),
        ("3. HYPOTHESES", s6.get("hypotheses_text", "")),
        ("4. METHOD", s6.get("method_section", "")),
    ]

    for heading_text, body in sections_data:
        h = doc.add_paragraph()
        run = h.add_run(heading_text)
        run.font.size = Pt(14)
        run.font.color.rgb = INDIGO
        run.bold = True
        if body:
            for para in body.split("\n\n"):
                p = doc.add_paragraph(para.strip())
                p.style.font.size = Pt(11)
        doc.add_paragraph()

    # Theoretical Implications
    h = doc.add_paragraph()
    run = h.add_run("5. THEORETICAL IMPLICATIONS")
    run.font.size = Pt(14)
    run.font.color.rgb = INDIGO
    run.bold = True
    for imp in s6.get("theoretical_implications", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{imp.get('id', '')}. {imp.get('title', '')}")
        run.bold = True
        run.font.size = Pt(11)
        doc.add_paragraph(imp.get("text", ""))

    doc.add_paragraph()

    # Practical Implications
    h = doc.add_paragraph()
    run = h.add_run("6. PRACTICAL IMPLICATIONS")
    run.font.size = Pt(14)
    run.font.color.rgb = INDIGO
    run.bold = True
    for imp in s6.get("practical_implications", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{imp.get('id', '')}. {imp.get('title', '')}")
        run.bold = True
        run.font.size = Pt(11)
        doc.add_paragraph(imp.get("text", ""))

    # Target journal note
    doc.add_paragraph()
    p = doc.add_paragraph()
    target = canvas["stage_5_review"].get("target_journal", "")
    run = p.add_run(f"Target Journal: {target}")
    run.italic = True
    run.font.color.rgb = INDIGO

    # Generated by
    p = doc.add_paragraph()
    run = p.add_run("Generated by Eagle — Research Copilot for GLIM Faculty")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

