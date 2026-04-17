"""
EAGLE — Research Copilot for GLIM Gurgaon Faculty
v2 — Editorial Redesign

Design system:
  - Editorial palette: off-white #FAFAF7 base, near-black #0A0A0A text,
    deep burgundy #7B1E3A single accent
  - Type pairing: Fraunces (serif display) + Inter (sans-serif UI/body)
  - Hairline borders, no gradients/blur/shadows, zero emoji decoration
  - Sidebar stepper (progress spine), persistent canvas rail, Cmd-K palette

Seven stages:
  1. Problematize    — Socratic → sharpened RQ
  2. Theorize        — construct map + hypothesis diagram (NEW)
  3. Literature      — OpenAlex-grounded lit review
  4. Method          — quant/qual design + interactive power calculator (NEW)
  5. Reviewer        — devil's advocate critique
  6. Full Paper      — auto-fix + manuscript generation
  7. Response Letter — R&R response letter generator (NEW)

Run:
    pip install -r requirements.txt
    streamlit run eagle_app.py
"""

import json
import sqlite3
import uuid
import datetime as dt
import io
import os
import re
from typing import Any

import streamlit as st
import streamlit.components.v1 as components
import requests
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============================================================================
# CONFIG
# =============================================================================

MODEL_HEAVY = "claude-3-sonnet-20240229"      # Deep reasoning
MODEL_LIGHT = "claude-3-haiku-20240307"        # Formatting, quick tasks
DB_PATH     = "eagle.db"
OPENALEX    = "https://api.openalex.org/works"

STAGES = [
    ("stage_1_problematization", "Problematize",   "01"),
    ("stage_2_theorization",     "Theorize",       "02"),
    ("stage_3_literature",       "Literature",     "03"),
    ("stage_4_method",           "Method",         "04"),
    ("stage_5_review",           "Reviewer",       "05"),
    ("stage_6_full_paper",       "Manuscript",     "06"),
    ("stage_7_response_letter",  "Response Letter","07"),
]

st.set_page_config(
    page_title="Eagle — Research Copilot",
    layout="wide",
    page_icon="◆",
    initial_sidebar_state="expanded",
)

# =============================================================================
# EDITORIAL DESIGN SYSTEM — CSS
# =============================================================================

DESIGN_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Fraunces:ital,opsz,wght@0,9..144,300;0,9..144,400;0,9..144,500;0,9..144,600;0,9..144,700;1,9..144,400&family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
  --bg: #FAFAF7;
  --bg-elevated: #FFFFFF;
  --ink: #0A0A0A;
  --ink-soft: #3F3F3E;
  --ink-muted: #787776;
  --ink-faint: #B4B3B0;
  --rule: rgba(10, 10, 10, 0.08);
  --rule-strong: rgba(10, 10, 10, 0.14);
  --accent: #7B1E3A;
  --accent-hover: #5E1129;
  --accent-soft: rgba(123, 30, 58, 0.06);
  --ok: #1F4A3C;
  --warn: #8B5A00;
  --danger: #8B1A1A;
  --serif: 'Fraunces', Georgia, serif;
  --sans: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
  --mono: 'JetBrains Mono', 'SF Mono', monospace;
}

/* ─── reset streamlit chrome ─── */
#MainMenu, footer, header, .stDeployButton { display: none !important; }
.stApp { background: var(--bg) !important; }
.block-container {
  padding-top: 2rem !important;
  padding-bottom: 6rem !important;
  max-width: 1400px !important;
}

/* ─── base type ─── */
html, body, [class*="css"], .stMarkdown, .stTextInput, .stTextArea,
.stSelectbox, .stRadio, .stButton, p, span, div, label {
  font-family: var(--sans) !important;
  color: var(--ink);
  letter-spacing: -0.005em;
}
.stMarkdown p {
  font-size: 15px;
  line-height: 1.65;
  color: var(--ink-soft);
}
h1, h2, h3, h4, h5 {
  font-family: var(--serif) !important;
  color: var(--ink) !important;
  letter-spacing: -0.02em;
  font-weight: 500 !important;
}
h1 { font-size: 2.5rem !important; line-height: 1.1 !important; }
h2 { font-size: 1.75rem !important; line-height: 1.2 !important; margin-top: 1.5rem !important; }
h3 { font-size: 1.25rem !important; line-height: 1.3 !important; }

/* ─── editorial wordmark ─── */
.eagle-wordmark {
  font-family: var(--serif);
  font-weight: 500;
  font-size: 24px;
  letter-spacing: -0.04em;
  color: var(--ink);
  display: flex;
  align-items: baseline;
  gap: 8px;
}
.eagle-wordmark .glyph {
  display: inline-block;
  width: 18px;
  height: 18px;
  background: var(--accent);
  transform: rotate(45deg);
  position: relative;
  top: 2px;
  margin-right: 2px;
}
.eagle-wordmark .sub {
  font-family: var(--sans);
  font-size: 10px;
  font-weight: 600;
  letter-spacing: 0.15em;
  color: var(--ink-muted);
  text-transform: uppercase;
  margin-left: 6px;
  align-self: center;
}

/* ─── sidebar ─── */
section[data-testid="stSidebar"] {
  background: var(--bg) !important;
  border-right: 1px solid var(--rule) !important;
  padding-top: 1.5rem;
}
section[data-testid="stSidebar"] > div:first-child {
  padding-top: 0 !important;
}
section[data-testid="stSidebar"] hr { display: none; }

/* stepper */
.stepper {
  margin: 24px 0 24px 0;
  padding: 0;
}
.stepper-item {
  display: flex;
  align-items: flex-start;
  gap: 14px;
  padding: 10px 0;
  border-left: 1px solid var(--rule);
  padding-left: 16px;
  margin-left: 8px;
  position: relative;
  cursor: pointer;
  transition: all 0.15s ease;
}
.stepper-item:hover { background: var(--accent-soft); }
.stepper-item .num {
  font-family: var(--mono);
  font-size: 11px;
  color: var(--ink-faint);
  font-weight: 500;
  letter-spacing: 0.08em;
  margin-top: 2px;
}
.stepper-item .label {
  font-family: var(--sans);
  font-size: 14px;
  font-weight: 500;
  color: var(--ink-soft);
}
.stepper-item .dot {
  position: absolute;
  left: -4px;
  top: 14px;
  width: 7px;
  height: 7px;
  border-radius: 50%;
  background: var(--bg);
  border: 1px solid var(--rule-strong);
}
.stepper-item.done .dot { background: var(--accent); border-color: var(--accent); }
.stepper-item.current { border-left-color: var(--accent); }
.stepper-item.current .dot {
  background: var(--accent);
  border-color: var(--accent);
  box-shadow: 0 0 0 4px var(--accent-soft);
}
.stepper-item.current .label { color: var(--ink); font-weight: 600; }
.stepper-item.done .label { color: var(--ink-muted); }

.sidebar-label {
  font-family: var(--sans);
  font-size: 10px;
  font-weight: 600;
  letter-spacing: 0.12em;
  text-transform: uppercase;
  color: var(--ink-muted);
  margin: 20px 0 10px 0;
}

/* ─── inputs ─── */
.stTextInput > div > div > input,
.stTextArea textarea,
.stSelectbox > div > div,
.stMultiSelect > div > div {
  background: var(--bg-elevated) !important;
  border: 1px solid var(--rule-strong) !important;
  border-radius: 4px !important;
  color: var(--ink) !important;
  font-family: var(--sans) !important;
  font-size: 14px !important;
  padding: 10px 12px !important;
  box-shadow: none !important;
  transition: border-color 0.15s ease !important;
}
.stTextInput > div > div > input:focus,
.stTextArea textarea:focus {
  border-color: var(--ink) !important;
  outline: none !important;
  box-shadow: none !important;
}
.stTextInput label, .stTextArea label, .stSelectbox label {
  font-family: var(--sans) !important;
  font-size: 12px !important;
  font-weight: 500 !important;
  color: var(--ink-muted) !important;
  letter-spacing: 0.02em !important;
  margin-bottom: 4px !important;
}

/* ─── buttons ─── */
.stButton > button {
  background: var(--ink) !important;
  color: var(--bg) !important;
  border: 1px solid var(--ink) !important;
  border-radius: 4px !important;
  font-family: var(--sans) !important;
  font-weight: 500 !important;
  font-size: 13px !important;
  padding: 9px 16px !important;
  letter-spacing: 0.01em !important;
  box-shadow: none !important;
  transition: all 0.12s ease !important;
  text-transform: none !important;
}
.stButton > button:hover {
  background: var(--accent) !important;
  border-color: var(--accent) !important;
  transform: none !important;
}
.stButton > button:focus { box-shadow: 0 0 0 3px var(--accent-soft) !important; }

/* secondary button styling — kind="secondary" */
.stButton > button[kind="secondary"] {
  background: var(--bg-elevated) !important;
  color: var(--ink) !important;
  border: 1px solid var(--rule-strong) !important;
}
.stButton > button[kind="secondary"]:hover {
  background: var(--bg) !important;
  border-color: var(--ink) !important;
  color: var(--ink) !important;
}

.stDownloadButton > button {
  background: var(--bg-elevated) !important;
  color: var(--ink) !important;
  border: 1px solid var(--ink) !important;
  border-radius: 4px !important;
  font-weight: 500 !important;
  font-size: 13px !important;
  padding: 9px 16px !important;
}
.stDownloadButton > button:hover {
  background: var(--ink) !important;
  color: var(--bg) !important;
}

/* ─── cards and surfaces ─── */
.eagle-card {
  background: var(--bg-elevated);
  border: 1px solid var(--rule);
  border-radius: 6px;
  padding: 24px;
  margin-bottom: 16px;
}
.eagle-card-bordered {
  border: 1px solid var(--rule-strong);
}

.eagle-kicker {
  font-family: var(--sans);
  font-size: 10px;
  font-weight: 600;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--accent);
  margin-bottom: 8px;
}
.eagle-display {
  font-family: var(--serif);
  font-size: 3rem;
  line-height: 1.05;
  letter-spacing: -0.03em;
  font-weight: 400;
  color: var(--ink);
  margin: 0 0 16px 0;
}
.eagle-lede {
  font-family: var(--serif);
  font-style: italic;
  font-size: 1.15rem;
  line-height: 1.5;
  color: var(--ink-soft);
  font-weight: 400;
}

/* ─── tabs stripped (we use sidebar stepper instead) ─── */
.stTabs [data-baseweb="tab-list"] {
  gap: 0 !important;
  border-bottom: 1px solid var(--rule) !important;
  background: transparent !important;
  padding: 0 !important;
  margin-bottom: 24px !important;
}
.stTabs [data-baseweb="tab"] {
  background: transparent !important;
  border: none !important;
  border-bottom: 2px solid transparent !important;
  border-radius: 0 !important;
  padding: 12px 18px !important;
  font-family: var(--sans) !important;
  font-size: 13px !important;
  font-weight: 500 !important;
  color: var(--ink-muted) !important;
  letter-spacing: 0.01em !important;
}
.stTabs [data-baseweb="tab"][aria-selected="true"] {
  color: var(--ink) !important;
  border-bottom-color: var(--ink) !important;
  background: transparent !important;
}

/* ─── expanders ─── */
.streamlit-expanderHeader, [data-testid="stExpander"] summary {
  background: var(--bg-elevated) !important;
  border: 1px solid var(--rule) !important;
  border-radius: 4px !important;
  font-family: var(--sans) !important;
  font-size: 13px !important;
  font-weight: 500 !important;
  color: var(--ink) !important;
  padding: 12px 16px !important;
}
[data-testid="stExpander"] {
  border: none !important;
  background: transparent !important;
}

/* ─── radio / select visuals ─── */
.stRadio > div { gap: 6px; }
.stRadio label {
  font-family: var(--sans) !important;
  font-size: 14px !important;
  color: var(--ink-soft) !important;
}

/* ─── alerts / info ─── */
.stAlert {
  border-radius: 4px !important;
  border: 1px solid var(--rule-strong) !important;
  background: var(--bg-elevated) !important;
  font-family: var(--sans) !important;
  font-size: 13px !important;
}

/* ─── metrics ─── */
[data-testid="stMetricValue"] {
  font-family: var(--serif) !important;
  font-weight: 500 !important;
  color: var(--ink) !important;
}
[data-testid="stMetricLabel"] {
  font-family: var(--sans) !important;
  font-size: 11px !important;
  font-weight: 600 !important;
  letter-spacing: 0.08em !important;
  text-transform: uppercase !important;
  color: var(--ink-muted) !important;
}

/* ─── dataframe ─── */
.stDataFrame {
  border: 1px solid var(--rule) !important;
  border-radius: 4px !important;
}

/* ─── rail (right side canvas summary) ─── */
.eagle-rail {
  position: sticky;
  top: 1rem;
  padding: 20px;
  background: var(--bg-elevated);
  border: 1px solid var(--rule);
  border-radius: 6px;
  font-size: 13px;
}
.eagle-rail h5 {
  font-family: var(--sans) !important;
  font-size: 10px !important;
  font-weight: 600 !important;
  letter-spacing: 0.14em !important;
  text-transform: uppercase !important;
  color: var(--ink-muted) !important;
  margin: 16px 0 6px 0 !important;
}
.eagle-rail .val {
  font-family: var(--serif);
  color: var(--ink);
  font-size: 14px;
  line-height: 1.45;
  font-style: italic;
}
.eagle-rail .empty {
  color: var(--ink-faint);
  font-style: italic;
  font-size: 12px;
}

/* ─── dividers ─── */
hr {
  border: none !important;
  border-top: 1px solid var(--rule) !important;
  margin: 24px 0 !important;
}

/* ─── chips ─── */
.chip {
  display: inline-block;
  padding: 3px 10px;
  border: 1px solid var(--rule-strong);
  border-radius: 999px;
  font-family: var(--sans);
  font-size: 11px;
  font-weight: 500;
  color: var(--ink-soft);
  margin-right: 6px;
  margin-bottom: 6px;
  background: var(--bg-elevated);
}
.chip.accent {
  background: var(--accent);
  color: white;
  border-color: var(--accent);
}
.chip.ok { border-color: var(--ok); color: var(--ok); }
.chip.warn { border-color: var(--warn); color: var(--warn); }
.chip.danger { border-color: var(--danger); color: var(--danger); }

/* ─── severity indicators ─── */
.sev-high { color: var(--danger); font-weight: 600; }
.sev-med { color: var(--warn); font-weight: 600; }
.sev-low { color: var(--ok); font-weight: 600; }

/* page header */
.eagle-page-head {
  display: flex;
  align-items: baseline;
  justify-content: space-between;
  margin-bottom: 8px;
  padding-bottom: 0;
}
.eagle-page-head .stage-num {
  font-family: var(--mono);
  font-size: 11px;
  color: var(--ink-faint);
  letter-spacing: 0.1em;
}

/* ─── progress ─── */
[data-testid="stProgressBar"] > div { background: var(--accent) !important; }

/* ─── hide label visibility when requested still needs height clean ─── */
.stTextArea [data-baseweb="textarea"] { min-height: 80px; }

/* ─── landing hero ─── */
.landing {
  max-width: 780px;
  margin: 4rem auto 2rem auto;
  padding: 0 1rem;
}
.landing .lede-big {
  font-family: var(--serif);
  font-size: 4rem;
  line-height: 1.02;
  letter-spacing: -0.035em;
  color: var(--ink);
  font-weight: 400;
  margin-bottom: 1.5rem;
}
.landing .lede-big em {
  font-style: italic;
  color: var(--accent);
}
.landing .sub {
  font-size: 1.15rem;
  line-height: 1.5;
  color: var(--ink-soft);
  max-width: 600px;
  margin-bottom: 3rem;
}
.landing .stages-grid {
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 1px;
  background: var(--rule);
  border: 1px solid var(--rule);
  border-radius: 4px;
  overflow: hidden;
  margin-bottom: 3rem;
}
.landing .stage-cell {
  background: var(--bg-elevated);
  padding: 20px 24px;
}
.landing .stage-cell .n {
  font-family: var(--mono);
  font-size: 11px;
  color: var(--ink-faint);
  letter-spacing: 0.1em;
  margin-bottom: 6px;
}
.landing .stage-cell .h {
  font-family: var(--serif);
  font-size: 18px;
  color: var(--ink);
  margin-bottom: 4px;
}
.landing .stage-cell .d {
  font-family: var(--sans);
  font-size: 13px;
  color: var(--ink-muted);
  line-height: 1.4;
}

/* code blocks (for system prompt previews etc) */
code, pre {
  font-family: var(--mono) !important;
  font-size: 12px !important;
}
</style>
"""

st.markdown(DESIGN_CSS, unsafe_allow_html=True)

# =============================================================================
# DATABASE
# =============================================================================

def db_init():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS sessions (
            session_id TEXT PRIMARY KEY,
            faculty_id TEXT,
            title TEXT,
            canvas_json TEXT,
            created_at TEXT,
            updated_at TEXT
        )
    """)
    conn.commit()
    conn.close()

def db_save(canvas: dict):
    canvas["updated_at"] = dt.datetime.utcnow().isoformat()
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        INSERT INTO sessions (session_id, faculty_id, title, canvas_json, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?)
        ON CONFLICT(session_id) DO UPDATE SET
            canvas_json=excluded.canvas_json,
            updated_at=excluded.updated_at,
            title=excluded.title
    """, (
        canvas["session_id"],
        canvas["faculty_id"],
        canvas.get("title", "Untitled"),
        json.dumps(canvas),
        canvas["created_at"],
        canvas["updated_at"],
    ))
    conn.commit()
    conn.close()

def db_list(faculty_id: str):
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(
        "SELECT session_id, title, updated_at FROM sessions WHERE faculty_id=? ORDER BY updated_at DESC",
        (faculty_id,)
    ).fetchall()
    conn.close()
    return rows

def db_load(session_id: str) -> dict | None:
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute("SELECT canvas_json FROM sessions WHERE session_id=?", (session_id,)).fetchone()
    conn.close()
    return json.loads(row[0]) if row else None

def db_delete(session_id: str):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("DELETE FROM sessions WHERE session_id=?", (session_id,))
    conn.commit()
    conn.close()

# =============================================================================
# CANVAS SCHEMA
# =============================================================================

def new_canvas(faculty_id: str) -> dict:
    now = dt.datetime.utcnow().isoformat()
    return {
        "session_id": str(uuid.uuid4()),
        "faculty_id": faculty_id,
        "title": "Untitled research session",
        "created_at": now,
        "updated_at": now,
        "current_stage": "stage_1_problematization",
        "stage_1_problematization": {
            "fuzzy_idea": "",
            "socratic_exchange": [],
            "candidate_rqs": [],
            "selected_rq": "",
            "faculty_notes": "",
        },
        "stage_2_theorization": {
            "candidate_theories": [],
            "selected_theory": "",
            "recommendation_rationale": "",
            "construct_map": {},
            "hypotheses": [],
            "faculty_notes": "",
        },
        "stage_3_literature": {
            "search_queries_used": [],
            "lit_table": [],
            "synthesis": "",
            "identified_gap": "",
            "faculty_notes": "",
        },
        "stage_4_method": {
            "research_approach": "quantitative",
            "recommended_method": "",
            "method_rationale": "",
            "sample_plan": {},
            "instrument": {"blocks": [], "attention_checks": [], "demographics": []},
            "interview_protocol": {"sections": [], "probes": [], "sampling_strategy": ""},
            "ethics_flags": [],
            "power_calc": {"effect_f2": 0.15, "alpha": 0.05, "power": 0.80,
                          "n_predictors": 5, "computed_n": 92},
            "faculty_notes": "",
        },
        "stage_5_review": {
            "reviewer_critique": [],
            "overall_readiness_score": 0,
            "target_journal": "JAMS",
            "journal_specific_notes": "",
            "faculty_notes": "",
        },
        "stage_6_full_paper": {
            "introduction": "",
            "literature_review": "",
            "hypotheses_text": "",
            "method_section": "",
            "theoretical_implications": [],
            "practical_implications": [],
            "revised_notes": "",
            "generated": False,
        },
        "stage_7_response_letter": {
            "reviewer_comments_raw": "",
            "parsed_comments": [],
            "response_letter": "",
            "revision_summary": "",
            "generated": False,
        },
    }

def ensure_canvas_shape(canvas: dict) -> dict:
    """Back-fill missing keys so old sessions still work."""
    defaults = new_canvas(canvas.get("faculty_id", ""))
    for k, v in defaults.items():
        if k not in canvas:
            canvas[k] = v
        elif isinstance(v, dict) and isinstance(canvas.get(k), dict):
            for subk, subv in v.items():
                if subk not in canvas[k]:
                    canvas[k][subk] = subv
    return canvas

def stage_completion(canvas: dict) -> dict[str, bool]:
    """Check which stages have meaningful output."""
    return {
        "stage_1_problematization": bool(canvas["stage_1_problematization"].get("selected_rq")),
        "stage_2_theorization": bool(canvas["stage_2_theorization"].get("construct_map")),
        "stage_3_literature": bool(canvas["stage_3_literature"].get("lit_table")),
        "stage_4_method": bool(
            canvas["stage_4_method"].get("instrument", {}).get("blocks") or
            canvas["stage_4_method"].get("interview_protocol", {}).get("sections")
        ),
        "stage_5_review": bool(canvas["stage_5_review"].get("reviewer_critique")),
        "stage_6_full_paper": bool(canvas["stage_6_full_paper"].get("generated")),
        "stage_7_response_letter": bool(canvas["stage_7_response_letter"].get("generated")),
    }

# =============================================================================
# OPENALEX
# =============================================================================

def openalex_search(query: str, per_page: int = 25) -> list[dict]:
    params = {
        "search": query,
        "per_page": per_page,
        "filter": "type:article,from_publication_date:2012-01-01",
        "sort": "cited_by_count:desc",
        "mailto": "research@glim.edu.in",
    }
    try:
        r = requests.get(OPENALEX, params=params, timeout=20)
        r.raise_for_status()
        data = r.json()
    except Exception as e:
        st.warning(f"OpenAlex error for '{query}': {e}")
        return []
    out = []
    for w in data.get("results", []):
        authors = ", ".join(a["author"]["display_name"]
                            for a in (w.get("authorships") or [])[:4])
        pl = w.get("primary_location") or {}
        src = pl.get("source") or {}
        out.append({
            "openalex_id": w.get("id", ""),
            "doi": w.get("doi", ""),
            "authors": authors,
            "year": w.get("publication_year", ""),
            "title": w.get("title", ""),
            "venue": src.get("display_name", "") or "",
            "abstract": reconstruct_abstract(w.get("abstract_inverted_index")),
            "cited_by": w.get("cited_by_count", 0),
        })
    return out

def reconstruct_abstract(inv_idx):
    if not inv_idx:
        return ""
    positions = []
    for word, idxs in inv_idx.items():
        for i in idxs:
            positions.append((i, word))
    positions.sort()
    return " ".join(w for _, w in positions)

# =============================================================================
# CLAUDE
# =============================================================================

def get_client() -> Anthropic:
    key = st.session_state.get("anthropic_api_key") or os.environ.get("ANTHROPIC_API_KEY")
    if not key:
        st.error("Enter your Anthropic API key in the sidebar to continue.")
        st.stop()
    return Anthropic(api_key=key.strip())

def claude_call(system: str, user: str, model: str = MODEL_HEAVY, max_tokens: int = 4000) -> str:
    custom_model = st.session_state.get("custom_model")
    if custom_model and custom_model.strip():
        model = custom_model.strip()
    client = get_client()
    try:
        resp = client.messages.create(
            model=model,
            system=system,
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": user}],
        )
        return resp.content[0].text
    except Exception as e:
        st.error(f"**Anthropic API Error:** {str(e)}")
        st.info(f"Model ID used: {model}")
        st.stop()

def claude_json(system: str, user: str, model: str = MODEL_HEAVY, max_tokens: int = 8000) -> dict:
    raw = claude_call(system, user, model=model, max_tokens=max_tokens)
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.split("```", 2)[1]
        if cleaned.startswith("json"):
            cleaned = cleaned[4:]
        cleaned = cleaned.rsplit("```", 1)[0]
    cleaned = cleaned.strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        # try to find a JSON object inside
        m = re.search(r'\{.*\}', cleaned, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                pass
        st.error("Couldn't parse JSON from the model. Raw output below. Try re-running.")
        st.code(raw[:3000])
        raise

# =============================================================================
# AGENT PROMPTS
# =============================================================================

SYS_PROBLEMATIZER_QUESTION = """You are the Problematizer agent in Eagle, a research copilot for marketing
and management faculty at a top Indian business school (GLIM Gurgaon).

You are talking to an experienced researcher. Match that register:
intellectually ambitious, concise, never condescending, no filler. Use
Alvesson & Sandberg's problematization method.

Your job RIGHT NOW: based on the conversation so far, ask ONE sharp Socratic
question that pushes the faculty toward clarity on the phenomenon, tension,
challenged assumption, stakeholders, or what would count as a surprising
finding. Never propose theories, methods, or citations. Output ONLY the
question text, no preamble, no quotes around it, no 'Q:' prefix.
"""

SYS_PROBLEMATIZER_FINALIZE = """You are the Problematizer agent. Based on the fuzzy idea and the Socratic
exchange below, produce THREE candidate research questions. Each must:
- Be empirically answerable
- Name the focal phenomenon
- Imply a tension or contrast
- Avoid generic "impact of X on Y" framing

Return ONLY valid JSON:
{"candidate_rqs": [{"rq": "...", "tradeoff": "..."}, ...]}
"""

SYS_THEORIST = """You are the Theorist agent in Eagle. The faculty has selected a research
question. Your job is to propose theoretical lenses and build a construct map
that a JAMS/JCR/JMR reviewer would find defensible.

Propose 2-3 candidate theories. For each: name, core claim, strengths,
weaknesses, 2-3 seminal references (only cite theories you are confident
exist). Recommend one. Build a construct map with focal IV(s), focal DV(s),
mediators, moderators, controls, boundary conditions. Derive 3-5 hypotheses.

Push back if the construct map reduces to "A affects B" — top-tier papers
need mechanism.

Return ONLY valid JSON:
{
  "candidate_theories": [{"name":"...","core_claim":"...","strengths":"...",
    "weaknesses":"...","seminal_refs":["..."]}],
  "recommended_theory": "...",
  "recommendation_rationale": "...",
  "construct_map": {
    "focal_iv": [{"name":"...","definition":"..."}],
    "focal_dv": [{"name":"...","definition":"..."}],
    "mediators": [{"name":"...","rationale":"..."}],
    "moderators": [{"name":"...","rationale":"...","moderates":"IV->Mediator or Mediator->DV or IV->DV"}],
    "controls": ["..."],
    "boundary_conditions": "..."
  },
  "hypotheses": [{"id":"H1","text":"...","type":"direct|mediation|moderation|moderated_mediation"}]
}
"""

SYS_LITSCOUT_QUERIES = """You are the Literature Scout. Given the construct map and theory below,
generate 5 targeted OpenAlex search queries. Each query should be 3-6 words,
combining constructs, theory, and context. Diversify — don't repeat the same
core phrase.

Return ONLY valid JSON: {"queries": ["...", "...", "...", "...", "..."]}
"""

SYS_LITSCOUT_SYNTHESIZE = """You are the Literature Scout. You have been given REAL papers retrieved
from OpenAlex. Score each 0-3 on relevance, keep the top ~15, build a
structured lit table, and write a 350-450 word synthesis identifying 2-3
specific gaps.

HARD RULES:
- Use ONLY the papers provided. Never invent DOIs, authors, or years.
- Never quote more than 10 words from any abstract. Paraphrase findings.
- If the abstract is empty, put "requires full-text review" in findings/gap.

Return ONLY valid JSON:
{
  "lit_table": [{"authors":"...","year":2020,"venue":"...","theory":"...",
    "context":"...","method":"...","key_constructs":"...","findings":"...",
    "gap_identified":"...","relevance_to_rq":"...","doi":"...","openalex_id":"..."}],
  "synthesis": "...",
  "identified_gap": "..."
}
"""

SYS_METHOD_QUANT = """You are the Method Designer agent in Eagle (quantitative track). Given
the construct map, hypotheses, and literature context, recommend a method
and generate a survey instrument.

For each construct, generate 4-5 items. Where a well-known validated scale
exists (e.g., Peck & Shu 2009 for psychological ownership, Brakus et al.
2009 for brand experience, Malhotra et al. 2006 for IUIPC), cite the source
and adapt. Otherwise generate new items following construct-definition-first
logic and flag as "newly generated — requires pretest."

Include: 2 attention checks, reverse-coded items (>=20% per scale), CMB
marker (Simmering et al. 2015), demographics block. Never pad.

Return ONLY valid JSON:
{
  "recommended_method": "...",
  "method_rationale": "...",
  "sample_plan": {"target_n": 0, "power_analysis": "...", "recruitment": "...",
    "inclusion_criteria": "..."},
  "instrument": {
    "blocks": [{"block_name":"...","construct":"...","scale_source":"...",
      "validated": true, "items":[{"id":"X1","text":"...","scale":"7-point Likert",
      "reverse_coded": false}]}],
    "attention_checks":[{"id":"AC1","text":"..."}],
    "cmb_marker":{"construct":"...","items":[]},
    "demographics":["age","gender","..."]
  },
  "ethics_flags": ["..."]
}
"""

SYS_METHOD_QUAL = """You are the Method Designer agent in Eagle (qualitative track). Given the
research question, theory, and construct map, design a rigorous qualitative
methodology with a semi-structured interview protocol.

Include: method (phenomenological, grounded theory, case study, ethnography,
narrative), sampling strategy, protocol with 3-5 thematic sections each
having 4-6 open-ended questions with 2-3 probes each, opening ice-breaker,
closing reflection. No yes/no questions.

Return ONLY valid JSON:
{
  "recommended_method": "...",
  "method_rationale": "...",
  "sample_plan": {"strategy":"...","target_n":0,"saturation_plan":"...",
    "recruitment":"...","inclusion_criteria":"..."},
  "interview_protocol": {
    "sections":[{"section_name":"...","construct_explored":"...",
      "questions":[{"id":"Q1","text":"...","probes":["...","..."]}]}],
    "opening_question":"...","closing_question":"...",
    "estimated_duration_minutes": 45
  },
  "analysis_approach": {"method":"...","coding_strategy":"...",
    "trustworthiness":["...","..."]},
  "ethics_flags": ["..."]
}
"""

SYS_REVIEWER = """You are the Reviewer agent in Eagle. You simulate a demanding but fair
reviewer at the target journal. Read the ENTIRE canvas. Find the concerns
that would get this paper desk-rejected or sent for major revision.

Evaluate: (a) Theoretical contribution (b) Problematization (c) Identification
and causal inference (d) Measurement (e) Generalizability (f) Alternative
explanations.

Generate 6-10 concerns. Be direct — soft feedback wastes time. Never
recommend rejection — recommend revision paths. Assign overall readiness
score 0-10 for the target journal.

Return ONLY valid JSON:
{
  "reviewer_critique": [{"concern_id":"R1","category":"...","severity":"low|medium|high",
    "issue":"...","suggested_fix":"...","linked_stage":"stage_2_theorization"}],
  "overall_readiness_score": 7.0,
  "journal_specific_notes": "..."
}
"""

SYS_FULL_PAPER = """You are the Paper Architect agent in Eagle. You have the ENTIRE research
canvas including reviewer critique. Auto-fix ALL reviewer concerns and
generate a publication-ready paper with fully written sections (APA style,
present tense for theory, past tense for methods).

Write for a top management/marketing journal. Each section should be
substantial (Introduction ~800 words, Literature Review ~1200 words,
Hypothesis Development ~1000 words, Method ~600 words).

For implications, give 4 theoretical contributions and 4 practical
implications, each 2-3 sentences substantive.

Return ONLY valid JSON:
{
  "introduction": "...",
  "literature_review": "...",
  "hypotheses_text": "...",
  "method_section": "...",
  "theoretical_implications": [{"title":"...","text":"..."}],
  "practical_implications": [{"title":"...","text":"..."}],
  "revised_notes": "How I addressed each reviewer concern: ..."
}
"""

SYS_PARSE_REVIEWER_COMMENTS = """You are a Response Letter assistant in Eagle. The faculty has pasted raw
reviewer comments from a real journal revision. Parse them into individual
numbered concerns.

Handle these common formats:
- "Reviewer 1: Comment 1: ..." style
- Numbered lists
- Prose paragraphs that mix multiple concerns (split them)
- Editor's letter + reviewer comments (tag each by source)

Return ONLY valid JSON:
{
  "parsed_comments": [
    {
      "id": "R1C1",
      "source": "Reviewer 1|Reviewer 2|Reviewer 3|Editor|AE",
      "comment_text": "verbatim extracted concern",
      "category": "theory|method|measurement|framing|writing|literature|contribution|other",
      "severity_hint": "minor|moderate|major"
    }
  ],
  "overall_tone": "brief read on whether reviewers were positive / mixed / hostile",
  "apparent_decision": "major revision | minor revision | conditional accept | split | reject-resubmit | unclear"
}
"""

SYS_RESPONSE_LETTER = """You are the Response Letter Architect in Eagle. You have (a) the full
research canvas, (b) the parsed reviewer comments, (c) the target journal.

Your job: draft a point-by-point response letter that editors love. For
each comment:
1. Restate the reviewer's concern in ONE sentence (respectful, not
   defensive)
2. Acknowledge the concern substantively (1-2 sentences) — never dismiss,
   even for ones you disagree with
3. Explain what was changed in the manuscript (2-4 sentences) with
   specific section/page references like "(Section 3.2, p. 14)"
4. Where the reviewer is wrong or asks for something inappropriate, push
   back politely with theoretical or empirical grounding

Use professional academic register. Open with gratitude. Close with a
summary of major changes. Tone: confident, collaborative, never obsequious.

Also produce a one-paragraph "revision summary" suitable for the cover
letter to the editor.

Return ONLY valid JSON:
{
  "response_letter": "full response letter as ONE markdown string with clear Reviewer 1 / Reviewer 2 / Editor headings and numbered responses",
  "revision_summary": "one paragraph (150-200 words) summarizing the revision for the editor"
}
"""

# =============================================================================
# FEATURE 1 — HYPOTHESIS PATH DIAGRAM
# =============================================================================

def render_path_diagram(construct_map: dict, hypotheses: list) -> str:
    """Auto-layout SEM-style path diagram as SVG. Returns HTML string for
    components.html. Layout: IVs on left, mediators center, DVs right,
    moderators as branching arrows with rotated boxes."""
    if not construct_map:
        return ""

    ivs = construct_map.get("focal_iv", []) or []
    dvs = construct_map.get("focal_dv", []) or []
    meds = construct_map.get("mediators", []) or []
    mods = construct_map.get("moderators", []) or []

    # canvas dimensions
    W, H = 1100, max(500, 140 + max(len(ivs), len(dvs), len(meds), 1) * 110)
    col_iv, col_med, col_dv = 120, W/2, W - 120

    def y_positions(n, top=100, bottom=H-80):
        if n == 0: return []
        if n == 1: return [(top + bottom) / 2]
        step = (bottom - top) / max(1, n - 1)
        return [top + i * step for i in range(n)]

    iv_ys  = y_positions(len(ivs))
    med_ys = y_positions(len(meds))
    dv_ys  = y_positions(len(dvs))

    # helper to escape
    def esc(s): return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;").replace('"',"&quot;")

    def box(x, y, w, h, label, kind="iv", hnum=""):
        if kind == "iv":
            fill, stroke, text_color = "#FFFFFF", "#0A0A0A", "#0A0A0A"
        elif kind == "dv":
            fill, stroke, text_color = "#7B1E3A", "#7B1E3A", "#FFFFFF"
        elif kind == "med":
            fill, stroke, text_color = "#FAFAF7", "#0A0A0A", "#0A0A0A"
        else:  # moderator
            fill, stroke, text_color = "#FFFFFF", "#787776", "#3F3F3E"

        # word-wrap the label
        words = label.split()
        lines = []
        cur = []
        max_char = int(w / 7.5)
        for word in words:
            if len(" ".join(cur + [word])) > max_char and cur:
                lines.append(" ".join(cur))
                cur = [word]
            else:
                cur.append(word)
        if cur: lines.append(" ".join(cur))
        lines = lines[:3]
        line_count = len(lines) or 1

        # adjust h for long labels
        eff_h = max(h, 24 + line_count * 16)

        svg = f'''<rect x="{x - w/2}" y="{y - eff_h/2}" width="{w}" height="{eff_h}"
                    fill="{fill}" stroke="{stroke}" stroke-width="{1.5 if kind=='dv' else 1}"
                    rx="3"/>'''
        line_y = y - ((line_count - 1) * 16) / 2
        for line in lines:
            svg += f'''<text x="{x}" y="{line_y + 5}" text-anchor="middle"
                        font-family="Inter, sans-serif" font-size="13"
                        font-weight="{'600' if kind=='dv' else '500'}"
                        fill="{text_color}">{esc(line)}</text>'''
            line_y += 16
        return svg, eff_h

    # SVG defs (arrowhead)
    defs = '''<defs>
      <marker id="arrow" viewBox="0 0 10 10" refX="9" refY="5" markerUnits="strokeWidth"
              markerWidth="8" markerHeight="8" orient="auto">
        <path d="M 0 0 L 10 5 L 0 10 z" fill="#0A0A0A"/>
      </marker>
      <marker id="arrow-accent" viewBox="0 0 10 10" refX="9" refY="5" markerUnits="strokeWidth"
              markerWidth="8" markerHeight="8" orient="auto">
        <path d="M 0 0 L 10 5 L 0 10 z" fill="#7B1E3A"/>
      </marker>
      <marker id="arrow-dash" viewBox="0 0 10 10" refX="9" refY="5" markerUnits="strokeWidth"
              markerWidth="7" markerHeight="7" orient="auto">
        <path d="M 0 0 L 10 5 L 0 10 z" fill="#787776"/>
      </marker>
    </defs>'''

    # build boxes
    nodes_svg = ""
    iv_positions = []
    for iv, y in zip(ivs, iv_ys):
        s, _ = box(col_iv, y, 180, 54, iv.get("name",""), "iv")
        nodes_svg += s
        iv_positions.append((col_iv + 90, y))  # right edge

    med_positions = []
    for med, y in zip(meds, med_ys):
        s, _ = box(col_med, y, 200, 54, med.get("name",""), "med")
        nodes_svg += s
        med_positions.append((col_med, y))  # center

    dv_positions = []
    for dv, y in zip(dvs, dv_ys):
        s, _ = box(col_dv, y, 180, 60, dv.get("name",""), "dv")
        nodes_svg += s
        dv_positions.append((col_dv - 90, y))  # left edge

    # draw paths
    # hypotheses are annotated with H# labels
    hypo_counter = iter([h.get("id", f"H{i+1}") for i, h in enumerate(hypotheses)])

    paths_svg = ""
    # IV -> Mediator (if mediators) then Mediator -> DV, else IV -> DV
    if meds:
        for ix, iy in iv_positions:
            for mx, my in med_positions:
                try: h_label = next(hypo_counter)
                except StopIteration: h_label = ""
                paths_svg += f'''<line x1="{ix}" y1="{iy}" x2="{mx - 100}" y2="{my}"
                  stroke="#0A0A0A" stroke-width="1.5" marker-end="url(#arrow)"/>'''
                if h_label:
                    paths_svg += f'''<text x="{(ix + mx - 100)/2}" y="{(iy + my)/2 - 6}"
                      font-family="JetBrains Mono, monospace" font-size="11" fill="#7B1E3A"
                      text-anchor="middle" font-weight="500">{esc(h_label)}</text>'''
        for mx, my in med_positions:
            for dx, dy in dv_positions:
                try: h_label = next(hypo_counter)
                except StopIteration: h_label = ""
                paths_svg += f'''<line x1="{mx + 100}" y1="{my}" x2="{dx}" y2="{dy}"
                  stroke="#0A0A0A" stroke-width="1.5" marker-end="url(#arrow)"/>'''
                if h_label:
                    paths_svg += f'''<text x="{(mx + 100 + dx)/2}" y="{(my + dy)/2 - 6}"
                      font-family="JetBrains Mono, monospace" font-size="11" fill="#7B1E3A"
                      text-anchor="middle" font-weight="500">{esc(h_label)}</text>'''
    else:
        # direct IV -> DV
        for ix, iy in iv_positions:
            for dx, dy in dv_positions:
                try: h_label = next(hypo_counter)
                except StopIteration: h_label = ""
                paths_svg += f'''<line x1="{ix}" y1="{iy}" x2="{dx}" y2="{dy}"
                  stroke="#0A0A0A" stroke-width="1.5" marker-end="url(#arrow)"/>'''
                if h_label:
                    paths_svg += f'''<text x="{(ix + dx)/2}" y="{(iy + dy)/2 - 6}"
                      font-family="JetBrains Mono, monospace" font-size="11" fill="#7B1E3A"
                      text-anchor="middle" font-weight="500">{esc(h_label)}</text>'''

    # Moderators — show as rotated boxes above or below paths
    mod_svg = ""
    for i, mod in enumerate(mods):
        # position moderator above each column of paths
        mx = col_med
        my = 40 + i * 50  # stack them at top
        name = mod.get("name","")
        s, _ = box(mx, my, 170, 40, name, "mod")
        mod_svg += s
        # dashed arrow from mod down to the midline
        paths_svg += f'''<line x1="{mx}" y1="{my + 20}" x2="{mx}" y2="{H/2 - 30}"
          stroke="#787776" stroke-width="1" stroke-dasharray="4 3"
          marker-end="url(#arrow-dash)"/>'''

    # Controls listed at bottom
    ctrls = construct_map.get("controls", []) or []
    ctrl_svg = ""
    if ctrls:
        ctrl_text = "Controls: " + ", ".join(ctrls)
        ctrl_svg = f'''<text x="{W/2}" y="{H - 30}" text-anchor="middle"
          font-family="Inter, sans-serif" font-size="11" fill="#787776"
          font-style="italic">{esc(ctrl_text)}</text>'''

    # Legend / header labels
    header_svg = f'''
    <text x="{col_iv}" y="50" text-anchor="middle" font-family="Inter, sans-serif"
      font-size="10" font-weight="600" fill="#787776" letter-spacing="1.5">INDEPENDENT</text>
    <text x="{col_med}" y="50" text-anchor="middle" font-family="Inter, sans-serif"
      font-size="10" font-weight="600" fill="#787776" letter-spacing="1.5">MEDIATORS</text>
    <text x="{col_dv}" y="50" text-anchor="middle" font-family="Inter, sans-serif"
      font-size="10" font-weight="600" fill="#787776" letter-spacing="1.5">DEPENDENT</text>
    '''

    svg = f'''<svg viewBox="0 0 {W} {H}" xmlns="http://www.w3.org/2000/svg"
                   style="width:100%;height:auto;background:#FAFAF7;border:1px solid rgba(10,10,10,0.08);border-radius:6px;">
      {defs}
      {header_svg}
      {paths_svg}
      {nodes_svg}
      {mod_svg}
      {ctrl_svg}
    </svg>'''
    return svg

# =============================================================================
# FEATURE 2 — INTERACTIVE POWER CALCULATOR (HTML/JS widget)
# =============================================================================

POWER_WIDGET_HTML = """
<!DOCTYPE html>
<html>
<head>
<style>
  @import url('https://fonts.googleapis.com/css2?family=Fraunces:wght@400;500&family=Inter:wght@400;500;600&family=JetBrains+Mono&display=swap');
  * { box-sizing: border-box; margin: 0; padding: 0; font-family: 'Inter', sans-serif; }
  body { background: transparent; color: #0A0A0A; }
  .widget {
    background: #FFFFFF;
    border: 1px solid rgba(10,10,10,0.14);
    border-radius: 6px;
    padding: 24px;
    display: grid;
    grid-template-columns: 1fr 320px;
    gap: 28px;
    align-items: start;
  }
  .controls { }
  .kicker {
    font-size: 10px; font-weight: 600; letter-spacing: 0.18em;
    text-transform: uppercase; color: #7B1E3A; margin-bottom: 6px;
  }
  .title {
    font-family: 'Fraunces', serif;
    font-size: 20px; font-weight: 500; color: #0A0A0A;
    margin-bottom: 16px; letter-spacing: -0.02em;
  }
  .slider-row {
    display: grid;
    grid-template-columns: 160px 1fr 60px;
    align-items: center;
    gap: 12px;
    padding: 10px 0;
    border-bottom: 1px solid rgba(10,10,10,0.06);
  }
  .slider-row:last-child { border-bottom: none; }
  .slider-row label {
    font-size: 13px; color: #3F3F3E; font-weight: 500;
  }
  .slider-row input[type=range] {
    -webkit-appearance: none; width: 100%; height: 3px;
    background: rgba(10,10,10,0.12); border-radius: 2px; outline: none;
  }
  .slider-row input[type=range]::-webkit-slider-thumb {
    -webkit-appearance: none; appearance: none;
    width: 16px; height: 16px; border-radius: 50%;
    background: #7B1E3A; cursor: pointer; border: 2px solid #FFFFFF;
    box-shadow: 0 0 0 1px #7B1E3A;
  }
  .slider-row input[type=range]::-moz-range-thumb {
    width: 16px; height: 16px; border-radius: 50%;
    background: #7B1E3A; cursor: pointer; border: 2px solid #FFFFFF;
  }
  .slider-row .val {
    font-family: 'JetBrains Mono', monospace;
    font-size: 13px; color: #0A0A0A; text-align: right;
  }
  .select-row {
    display: grid; grid-template-columns: 160px 1fr;
    gap: 12px; padding: 10px 0;
    border-bottom: 1px solid rgba(10,10,10,0.06);
    align-items: center;
  }
  .select-row select {
    padding: 8px 10px; border: 1px solid rgba(10,10,10,0.14);
    border-radius: 4px; font-family: 'Inter', sans-serif;
    font-size: 13px; background: #FFFFFF; color: #0A0A0A;
  }
  .result {
    background: #FAFAF7;
    border: 1px solid rgba(10,10,10,0.08);
    border-radius: 6px;
    padding: 24px;
    text-align: center;
  }
  .result .label {
    font-size: 10px; font-weight: 600; letter-spacing: 0.14em;
    text-transform: uppercase; color: #787776; margin-bottom: 12px;
  }
  .result .n {
    font-family: 'Fraunces', serif;
    font-size: 72px; font-weight: 400; color: #0A0A0A;
    letter-spacing: -0.04em; line-height: 1;
  }
  .result .suffix {
    font-family: 'Inter', sans-serif; font-size: 13px;
    color: #787776; margin-top: 6px;
  }
  .result .formula {
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px; color: #B4B3B0;
    margin-top: 16px; padding-top: 12px;
    border-top: 1px solid rgba(10,10,10,0.06);
  }
  .interpret {
    font-family: 'Fraunces', serif; font-style: italic;
    font-size: 13px; color: #3F3F3E; margin-top: 14px;
    line-height: 1.4;
  }
  .effect-legend {
    display: flex; gap: 12px; font-size: 10px;
    color: #787776; margin-top: 4px;
  }
  .effect-legend span { font-family: 'JetBrains Mono', monospace; }
</style>
</head>
<body>
<div class="widget">
  <div class="controls">
    <div class="kicker">Power Analysis</div>
    <div class="title">G*Power — F-test, linear multiple regression</div>

    <div class="select-row">
      <label>Test type</label>
      <select id="testtype">
        <option value="regression">Multiple regression (F-test)</option>
        <option value="anova">One-way ANOVA (F-test)</option>
        <option value="interaction">Interaction term (f² for ΔR²)</option>
      </select>
    </div>

    <div class="slider-row">
      <label>Effect size (f²)</label>
      <input type="range" id="f2" min="0.02" max="0.50" step="0.01" value="0.15">
      <span class="val" id="f2val">0.15</span>
    </div>
    <div class="effect-legend" style="padding: 0 0 8px 172px;">
      <span>0.02 small</span><span>0.15 medium</span><span>0.35 large</span>
    </div>

    <div class="slider-row">
      <label>α (significance)</label>
      <input type="range" id="alpha" min="0.01" max="0.10" step="0.01" value="0.05">
      <span class="val" id="alphaval">0.05</span>
    </div>

    <div class="slider-row">
      <label>Power (1–β)</label>
      <input type="range" id="power" min="0.70" max="0.99" step="0.01" value="0.80">
      <span class="val" id="powerval">0.80</span>
    </div>

    <div class="slider-row">
      <label>Predictors</label>
      <input type="range" id="k" min="1" max="20" step="1" value="5">
      <span class="val" id="kval">5</span>
    </div>

    <div class="interpret" id="interpret"></div>
  </div>

  <div class="result">
    <div class="label">Required sample size</div>
    <div class="n" id="n_required">—</div>
    <div class="suffix">participants</div>
    <div class="formula" id="formula"></div>
  </div>
</div>

<script>
/* Approximation to G*Power F-test sample size calculation
   using Cohen's formula with noncentrality parameter λ = f² × (v + u + 1)
   where u = numerator df, v = denominator df.
   Iterative search for smallest n satisfying target power. */

function gpowerN(f2, alpha, power, k, testType) {
  // u = numerator df; varies by test
  let u;
  if (testType === 'regression') u = k;
  else if (testType === 'anova') u = k - 1; // groups - 1
  else if (testType === 'interaction') u = 1; // single interaction term
  if (u < 1) u = 1;

  // binary search for n
  // v = n - u - 1 (for regression); λ = f2 * (u + v + 1) = f2 * n
  // required: noncentral F critical > central F critical at given alpha/power
  // Shortcut: use Cohen's approximation L = f² × n; look up L for given u, power, alpha

  // L-table approximation from Cohen (1988) Table 9.4.2 — linear interp between key rows
  // keys: (u, alpha=0.05, power=0.80)
  const Ltable_p80_a05 = {1:7.85,2:9.64,3:10.90,4:11.94,5:12.83,6:13.62,7:14.35,8:15.02,10:16.25,15:18.88,20:21.20};
  // For adjustments with different alpha/power, use multiplicative factors
  // power adjustment
  const powerFactor = {0.70:0.85,0.75:0.92,0.80:1.00,0.85:1.10,0.90:1.22,0.95:1.40,0.99:1.75};
  const alphaFactor = {0.01:1.50,0.02:1.33,0.05:1.00,0.10:0.78};

  // nearest key for u
  const uKeys = Object.keys(Ltable_p80_a05).map(Number).sort((a,b)=>a-b);
  let uLo = uKeys[0], uHi = uKeys[uKeys.length-1];
  for (let i = 0; i < uKeys.length - 1; i++) {
    if (uKeys[i] <= u && u <= uKeys[i+1]) { uLo = uKeys[i]; uHi = uKeys[i+1]; break; }
  }
  if (u >= uHi) uLo = uHi;
  let Lbase;
  if (uLo === uHi) Lbase = Ltable_p80_a05[uLo];
  else {
    const frac = (u - uLo) / (uHi - uLo);
    Lbase = Ltable_p80_a05[uLo] + frac * (Ltable_p80_a05[uHi] - Ltable_p80_a05[uLo]);
  }

  // power multiplier (interp)
  function interpFactor(table, val) {
    const keys = Object.keys(table).map(Number).sort((a,b)=>a-b);
    if (val <= keys[0]) return table[keys[0]];
    if (val >= keys[keys.length-1]) return table[keys[keys.length-1]];
    for (let i = 0; i < keys.length - 1; i++) {
      if (keys[i] <= val && val <= keys[i+1]) {
        const f = (val - keys[i]) / (keys[i+1] - keys[i]);
        return table[keys[i]] + f * (table[keys[i+1]] - table[keys[i]]);
      }
    }
    return 1.0;
  }
  const pFac = interpFactor(powerFactor, power);
  const aFac = interpFactor(alphaFactor, alpha);
  const L = Lbase * pFac * aFac;

  let n = Math.ceil(L / f2) + u + 1;
  if (n < u + 3) n = u + 3;
  return n;
}

function update() {
  const f2 = parseFloat(document.getElementById('f2').value);
  const alpha = parseFloat(document.getElementById('alpha').value);
  const power = parseFloat(document.getElementById('power').value);
  const k = parseInt(document.getElementById('k').value);
  const testType = document.getElementById('testtype').value;

  document.getElementById('f2val').textContent = f2.toFixed(2);
  document.getElementById('alphaval').textContent = alpha.toFixed(2);
  document.getElementById('powerval').textContent = power.toFixed(2);
  document.getElementById('kval').textContent = k;

  const n = gpowerN(f2, alpha, power, k, testType);
  document.getElementById('n_required').textContent = n;

  let testLabel;
  if (testType === 'regression') testLabel = `regression with ${k} predictors`;
  else if (testType === 'anova') testLabel = `ANOVA with ${k} groups`;
  else testLabel = `interaction term, ΔR² via f²`;

  document.getElementById('formula').textContent =
    `F-test · u=${testType==='regression'?k:testType==='anova'?k-1:1} · λ = f² × n`;

  let effectDesc;
  if (f2 <= 0.05) effectDesc = "very small";
  else if (f2 <= 0.10) effectDesc = "small-to-medium";
  else if (f2 <= 0.20) effectDesc = "medium";
  else if (f2 <= 0.35) effectDesc = "medium-to-large";
  else effectDesc = "large";

  document.getElementById('interpret').innerHTML =
    `For a <b>${effectDesc}</b> effect (f²=${f2.toFixed(2)}) with ` +
    `${testLabel}, at α=${alpha.toFixed(2)} and power=${power.toFixed(2)}, ` +
    `you need <b>n=${n}</b>.`;

  // persist to Streamlit via postMessage
  window.parent.postMessage({
    type: 'eagle_power',
    f2: f2, alpha: alpha, power: power, k: k, n: n, testType: testType
  }, '*');
}

['f2','alpha','power','k','testtype'].forEach(id => {
  document.getElementById(id).addEventListener('input', update);
  document.getElementById(id).addEventListener('change', update);
});
update();
</script>
</body>
</html>
"""

# =============================================================================
# WORD EXPORT
# =============================================================================

INK = RGBColor(0x0A, 0x0A, 0x0A)
ACCENT = RGBColor(0x7B, 0x1E, 0x3A)
MUTED = RGBColor(0x78, 0x77, 0x76)

def _style_heading(p, color=INK, size=14, bold=True):
    for r in p.runs:
        r.font.name = "Georgia"
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.bold = bold

def export_canvas_docx(canvas: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)

    # title
    t = doc.add_heading(canvas.get("title","Research Canvas"), level=0)
    _style_heading(t, INK, 22)

    meta = doc.add_paragraph()
    mr = meta.add_run(f"Eagle Research Canvas · {canvas['session_id'][:8]} · Updated {canvas['updated_at'][:10]}")
    mr.font.size = Pt(9); mr.font.color.rgb = MUTED; mr.italic = True

    doc.add_paragraph()

    # Stage 1
    s1 = canvas["stage_1_problematization"]
    h = doc.add_heading("01. Problematization", level=1); _style_heading(h, ACCENT, 16)
    doc.add_paragraph(f"Fuzzy idea: {s1.get('fuzzy_idea','')}")
    p = doc.add_paragraph()
    r = p.add_run("Selected research question: "); r.bold = True
    p.add_run(s1.get("selected_rq",""))

    # Stage 2
    s2 = canvas["stage_2_theorization"]
    h = doc.add_heading("02. Theorization", level=1); _style_heading(h, ACCENT, 16)
    doc.add_paragraph(f"Theory: {s2.get('selected_theory','')}")
    if s2.get("recommendation_rationale"):
        p = doc.add_paragraph(); p.add_run("Rationale: ").bold=True; p.add_run(s2["recommendation_rationale"])

    cm = s2.get("construct_map") or {}
    if cm:
        doc.add_paragraph()
        for key, label in [("focal_iv","Independent"), ("focal_dv","Dependent"),
                           ("mediators","Mediators"), ("moderators","Moderators")]:
            items = cm.get(key, []) or []
            if items:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run("; ".join(i.get("name","") for i in items))
        if cm.get("controls"):
            p = doc.add_paragraph(); p.add_run("Controls: ").bold=True
            p.add_run(", ".join(cm["controls"]))
        if cm.get("boundary_conditions"):
            p = doc.add_paragraph(); p.add_run("Boundary conditions: ").bold=True
            p.add_run(cm["boundary_conditions"])
    for h_ in s2.get("hypotheses", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{h_.get('id','H')}: ").bold = True
        p.add_run(h_.get("text",""))

    # Stage 3
    s3 = canvas["stage_3_literature"]
    h = doc.add_heading("03. Literature Review", level=1); _style_heading(h, ACCENT, 16)
    lit = s3.get("lit_table") or []
    if lit:
        table = doc.add_table(rows=1, cols=5); table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, label in enumerate(["Authors (Year)","Venue","Theory / Method","Findings","Gap"]):
            hdr[i].text = label
            for r in hdr[i].paragraphs[0].runs: r.bold=True
        for p in lit:
            row = table.add_row().cells
            row[0].text = f"{p.get('authors','')} ({p.get('year','')})"
            row[1].text = p.get("venue","")
            row[2].text = f"{p.get('theory','')} · {p.get('method','')}"
            row[3].text = p.get("findings","")
            row[4].text = p.get("gap_identified","")
    if s3.get("synthesis"):
        doc.add_paragraph()
        h_ = doc.add_heading("Synthesis", level=2); _style_heading(h_, INK, 13)
        doc.add_paragraph(s3["synthesis"])
    if s3.get("identified_gap"):
        p = doc.add_paragraph()
        p.add_run("Identified gap: ").bold = True
        p.add_run(s3["identified_gap"])

    # Stage 4
    s4 = canvas["stage_4_method"]
    h = doc.add_heading("04. Method & Instrument", level=1); _style_heading(h, ACCENT, 16)
    doc.add_paragraph(f"Approach: {s4.get('research_approach','quantitative').title()}")
    doc.add_paragraph(f"Method: {s4.get('recommended_method','')}")
    doc.add_paragraph(f"Rationale: {s4.get('method_rationale','')}")
    sp = s4.get("sample_plan", {}) or {}
    if sp:
        p = doc.add_paragraph()
        p.add_run(f"Target N: {sp.get('target_n','')}  ·  Recruitment: {sp.get('recruitment','')}").italic = True
    pc = s4.get("power_calc", {}) or {}
    if pc and pc.get("computed_n"):
        p = doc.add_paragraph()
        p.add_run(f"Power analysis: f²={pc.get('effect_f2','')}, α={pc.get('alpha','')}, "
                  f"1−β={pc.get('power','')}, k={pc.get('n_predictors','')} → n={pc.get('computed_n','')}").italic=True

    # quantitative instrument
    instr = s4.get("instrument", {}) or {}
    for block in instr.get("blocks", []) or []:
        h_ = doc.add_heading(block.get("block_name","Block"), level=2); _style_heading(h_, INK, 13)
        p = doc.add_paragraph()
        p.add_run(f"Construct: {block.get('construct','')}  ·  "
                  f"Source: {block.get('scale_source','')}  ·  "
                  f"Validated: {block.get('validated','')}").italic=True
        for item in block.get("items", []) or []:
            tag = " (R)" if item.get("reverse_coded") else ""
            doc.add_paragraph(f"{item.get('id','')}{tag}: {item.get('text','')}", style="List Number")
    if instr.get("attention_checks"):
        h_ = doc.add_heading("Attention checks", level=2); _style_heading(h_, INK, 13)
        for ac in instr["attention_checks"]:
            doc.add_paragraph(f"{ac.get('id','')}: {ac.get('text','')}", style="List Bullet")

    # qualitative protocol
    proto = s4.get("interview_protocol") or {}
    if proto.get("sections"):
        h_ = doc.add_heading("Interview Protocol", level=2); _style_heading(h_, INK, 13)
        if proto.get("opening_question"):
            doc.add_paragraph(f"Opening: {proto['opening_question']}")
        for sec in proto.get("sections", []):
            h_ = doc.add_heading(sec.get("section_name",""), level=3); _style_heading(h_, INK, 12)
            for q in sec.get("questions", []):
                doc.add_paragraph(f"{q.get('id','')}: {q.get('text','')}", style="List Number")
                for probe in q.get("probes", []):
                    doc.add_paragraph(f"— {probe}", style="List Bullet 2")
        if proto.get("closing_question"):
            doc.add_paragraph(f"Closing: {proto['closing_question']}")

    # Stage 5
    s5 = canvas["stage_5_review"]
    if s5.get("reviewer_critique"):
        h = doc.add_heading("05. Reviewer Critique", level=1); _style_heading(h, ACCENT, 16)
        doc.add_paragraph(f"Target journal: {s5.get('target_journal','')}  ·  "
                          f"Readiness: {s5.get('overall_readiness_score','')}/10")
        for c in s5["reviewer_critique"]:
            p = doc.add_paragraph()
            p.add_run(f"[{c.get('severity','').upper()}] {c.get('category','')}: ").bold = True
            p.add_run(c.get("issue",""))
            doc.add_paragraph(f"   → Fix: {c.get('suggested_fix','')}")

    # Stage 6 — Manuscript
    s6 = canvas["stage_6_full_paper"]
    if s6.get("generated"):
        h = doc.add_heading("06. Manuscript Draft", level=1); _style_heading(h, ACCENT, 16)
        for key, label in [("introduction","Introduction"),
                            ("literature_review","Literature Review"),
                            ("hypotheses_text","Hypothesis Development"),
                            ("method_section","Method")]:
            if s6.get(key):
                h_ = doc.add_heading(label, level=2); _style_heading(h_, INK, 13)
                for para in str(s6[key]).split("\n\n"):
                    if para.strip(): doc.add_paragraph(para.strip())
        if s6.get("theoretical_implications"):
            h_ = doc.add_heading("Theoretical Implications", level=2); _style_heading(h_, INK, 13)
            for imp in s6["theoretical_implications"]:
                p = doc.add_paragraph()
                p.add_run(f"{imp.get('title','')}. ").bold = True
                p.add_run(imp.get("text",""))
        if s6.get("practical_implications"):
            h_ = doc.add_heading("Practical Implications", level=2); _style_heading(h_, INK, 13)
            for imp in s6["practical_implications"]:
                p = doc.add_paragraph()
                p.add_run(f"{imp.get('title','')}. ").bold = True
                p.add_run(imp.get("text",""))

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()

def export_response_letter_docx(canvas: dict) -> bytes:
    """Standalone response letter document — clean, journal-ready."""
    s7 = canvas["stage_7_response_letter"]
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)

    t = doc.add_heading("Response to Reviewers", level=0); _style_heading(t, INK, 20)
    meta = doc.add_paragraph()
    mr = meta.add_run(f"Manuscript: {canvas.get('title','')}")
    mr.font.size = Pt(10); mr.font.color.rgb = MUTED; mr.italic = True
    target = canvas["stage_5_review"].get("target_journal","")
    if target:
        mr2 = meta.add_run(f"   ·   Target journal: {target}")
        mr2.font.size = Pt(10); mr2.font.color.rgb = MUTED; mr2.italic = True

    doc.add_paragraph()

    if s7.get("revision_summary"):
        h = doc.add_heading("Summary of Revisions", level=1); _style_heading(h, ACCENT, 14)
        doc.add_paragraph(s7["revision_summary"])
        doc.add_paragraph()

    h = doc.add_heading("Point-by-Point Response", level=1); _style_heading(h, ACCENT, 14)
    letter = s7.get("response_letter","")
    for raw_para in letter.split("\n\n"):
        para = raw_para.strip()
        if not para: continue
        # markdown-ish parsing
        if para.startswith("## "):
            h_ = doc.add_heading(para[3:].strip(), level=2); _style_heading(h_, INK, 13)
        elif para.startswith("### "):
            h_ = doc.add_heading(para[4:].strip(), level=3); _style_heading(h_, INK, 12)
        elif para.startswith("**") and para.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(para.strip("*")).bold = True
        else:
            p = doc.add_paragraph(para)

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()

# =============================================================================
# UI — INITIALIZATION
# =============================================================================

db_init()
# migrate old db if present
if os.path.exists("sherpa.db") and not os.path.exists(DB_PATH):
    import shutil
    shutil.copy2("sherpa.db", DB_PATH)

# session state
if "canvas" not in st.session_state: st.session_state.canvas = None
if "faculty_id" not in st.session_state: st.session_state.faculty_id = ""
if "anthropic_api_key" not in st.session_state: st.session_state.anthropic_api_key = ""
if "current_stage_idx" not in st.session_state: st.session_state.current_stage_idx = 0

# =============================================================================
# SIDEBAR
# =============================================================================

with st.sidebar:
    # editorial wordmark
    st.markdown("""
    <div class="eagle-wordmark">
      <span class="glyph"></span>
      Eagle<span class="sub">GLIM · v2</span>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="sidebar-label">Credentials</div>', unsafe_allow_html=True)
    st.session_state.anthropic_api_key = st.text_input(
        "API key", type="password",
        value=st.session_state.anthropic_api_key,
        placeholder="API key (optional — leave blank for free tier)",
        label_visibility="collapsed",
    )
    st.session_state.faculty_id = st.text_input(
        "Faculty email",
        value=st.session_state.faculty_id or "harish@glim.ac.in",
        label_visibility="collapsed",
    )
    
    st.session_state.custom_model = st.text_input(
        "Model override",
        value=st.session_state.get("custom_model", ""),
        placeholder="Override Model ID (e.g. claude-3-haiku-20240307)",
        label_visibility="collapsed",
    )

    # Stepper
    if st.session_state.canvas:
        canvas = ensure_canvas_shape(st.session_state.canvas)
        st.session_state.canvas = canvas
        completion = stage_completion(canvas)
        st.markdown('<div class="sidebar-label">Progress</div>', unsafe_allow_html=True)

        stepper_html = '<div class="stepper">'
        for idx, (key, label, num) in enumerate(STAGES):
            done_class = "done" if completion.get(key) else ""
            current_class = "current" if idx == st.session_state.current_stage_idx else ""
            cls = f"stepper-item {done_class} {current_class}".strip()
            stepper_html += f'''
            <div class="{cls}" onclick="window.parent.postMessage({{type:'eagle_step', idx:{idx}}}, '*')">
              <div class="dot"></div>
              <div class="num">{num}</div>
              <div class="label">{label}</div>
            </div>
            '''
        stepper_html += '</div>'
        st.markdown(stepper_html, unsafe_allow_html=True)

        # fallback buttons (since onclick won't roundtrip to Streamlit without component)
        st.markdown('<div class="sidebar-label">Jump to stage</div>', unsafe_allow_html=True)
        for idx, (key, label, num) in enumerate(STAGES):
            is_current = idx == st.session_state.current_stage_idx
            marker = "●" if completion.get(key) else "○"
            btn_label = f"{num}  {marker}  {label}"
            if st.button(btn_label, key=f"nav_{idx}", use_container_width=True,
                         type="primary" if is_current else "secondary"):
                st.session_state.current_stage_idx = idx
                st.rerun()

    st.markdown("<hr/>", unsafe_allow_html=True)

    # Session controls
    st.markdown('<div class="sidebar-label">Sessions</div>', unsafe_allow_html=True)
    if st.button("+ New session", use_container_width=True, type="secondary"):
        st.session_state.canvas = new_canvas(st.session_state.faculty_id)
        db_save(st.session_state.canvas)
        st.session_state.current_stage_idx = 0
        st.rerun()

    if st.session_state.faculty_id:
        rows = db_list(st.session_state.faculty_id)
        if rows:
            st.markdown('<div class="sidebar-label" style="margin-top:16px;">Recent</div>', unsafe_allow_html=True)
            for sid, title, updated in rows[:8]:
                short_title = (title or "Untitled")[:28]
                if st.button(f"· {short_title}", key=f"load_{sid}",
                             use_container_width=True, type="secondary"):
                    st.session_state.canvas = ensure_canvas_shape(db_load(sid))
                    st.session_state.current_stage_idx = 0
                    st.rerun()

    # Export
    if st.session_state.canvas:
        st.markdown("<hr/>", unsafe_allow_html=True)
        st.markdown('<div class="sidebar-label">Export</div>', unsafe_allow_html=True)
        try:
            docx_bytes = export_canvas_docx(st.session_state.canvas)
            st.download_button(
                "Download canvas (.docx)",
                data=docx_bytes,
                file_name=f"eagle_canvas_{st.session_state.canvas['session_id'][:8]}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as e:
            st.caption(f"Export error: {e}")

# =============================================================================
# LANDING (no canvas)
# =============================================================================

if not st.session_state.canvas:
    st.markdown("""
    <div class="landing">
      <div class="lede-big">A research partner<br/>for <em>serious</em> faculty.</div>
      <div class="sub">Eagle walks you from a fuzzy idea to a publication-ready manuscript through seven deliberate stages. Built for top-tier journals, grounded in real literature, and designed to make the hardest parts of research — problematizing, theorizing, reviewing — tractable.</div>
      <div class="stages-grid">
        <div class="stage-cell"><div class="n">01</div><div class="h">Problematize</div><div class="d">Sharpen a fuzzy idea into a defensible research question via Socratic dialogue.</div></div>
        <div class="stage-cell"><div class="n">02</div><div class="h">Theorize</div><div class="d">Choose a lens, build a construct map, and auto-generate a hypothesis path diagram.</div></div>
        <div class="stage-cell"><div class="n">03</div><div class="h">Literature</div><div class="d">Real papers from OpenAlex, structured into a gap-finding table with synthesis.</div></div>
        <div class="stage-cell"><div class="n">04</div><div class="h">Method</div><div class="d">Quant or qual design, validated instruments, and an interactive G*Power calculator.</div></div>
        <div class="stage-cell"><div class="n">05</div><div class="h">Reviewer</div><div class="d">Devil's advocate critique calibrated to the target journal.</div></div>
        <div class="stage-cell"><div class="n">06</div><div class="h">Manuscript</div><div class="d">Full draft with Introduction, LR, Hypothesis Development, Method, and Implications.</div></div>
        <div class="stage-cell" style="grid-column: 1 / -1;"><div class="n">07</div><div class="h">Response Letter</div><div class="d">Paste reviewer comments from an R&R and Eagle drafts a point-by-point response letter in journal format.</div></div>
      </div>
    </div>
    """, unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c2:
        if st.button("Start a new session →", use_container_width=True):
            if not st.session_state.faculty_id:
                st.warning("Enter your email in the sidebar first.")
            else:
                st.session_state.canvas = new_canvas(st.session_state.faculty_id)
                db_save(st.session_state.canvas)
                st.rerun()
    st.stop()

canvas = ensure_canvas_shape(st.session_state.canvas)
st.session_state.canvas = canvas

# =============================================================================
# MAIN — two-column layout: stage work (left) + persistent canvas rail (right)
# =============================================================================

# session title bar (above everything)
title_col, meta_col = st.columns([5, 2])
with title_col:
    new_title = st.text_input("Session title", value=canvas.get("title","Untitled research session"),
                              label_visibility="collapsed")
    if new_title != canvas.get("title"):
        canvas["title"] = new_title
        db_save(canvas)
with meta_col:
    st.markdown(
        f'<div style="text-align:right;padding-top:8px;font-family:var(--mono);'
        f'font-size:11px;color:#B4B3B0;letter-spacing:0.08em;">'
        f'ID {canvas["session_id"][:8]} · {canvas["updated_at"][:10]}'
        f'</div>',
        unsafe_allow_html=True
    )

st.markdown("<hr/>", unsafe_allow_html=True)

# Main grid
main_col, rail_col = st.columns([3.2, 1.4], gap="large")

# -------- RIGHT RAIL (persistent canvas summary) --------
with rail_col:
    s1 = canvas["stage_1_problematization"]
    s2 = canvas["stage_2_theorization"]
    s3 = canvas["stage_3_literature"]
    s4 = canvas["stage_4_method"]
    s5 = canvas["stage_5_review"]

    rq_text = s1.get("selected_rq") or ""
    theory_text = s2.get("selected_theory") or ""
    gap_text = s3.get("identified_gap") or ""
    method_text = s4.get("recommended_method") or ""
    journal = s5.get("target_journal") or ""
    score = s5.get("overall_readiness_score") or 0

    rail_html = '<div class="eagle-rail">'
    rail_html += '<div style="font-family:var(--sans);font-size:10px;font-weight:600;letter-spacing:0.14em;text-transform:uppercase;color:#7B1E3A;">Research Canvas</div>'

    rail_html += '<h5>Research Question</h5>'
    if rq_text:
        rail_html += f'<div class="val">{rq_text[:260]}{"…" if len(rq_text)>260 else ""}</div>'
    else:
        rail_html += '<div class="empty">Complete Stage 01</div>'

    rail_html += '<h5>Theoretical Lens</h5>'
    if theory_text:
        rail_html += f'<div class="val">{theory_text[:200]}</div>'
    else:
        rail_html += '<div class="empty">Complete Stage 02</div>'

    cm = s2.get("construct_map") or {}
    if cm:
        ivs = [i.get("name","") for i in cm.get("focal_iv",[])]
        dvs = [i.get("name","") for i in cm.get("focal_dv",[])]
        meds = [i.get("name","") for i in cm.get("mediators",[])]
        rail_html += '<h5>Constructs</h5>'
        rail_html += '<div style="font-family:var(--sans);font-size:12px;color:#3F3F3E;">'
        if ivs: rail_html += f'<div style="margin-bottom:4px;"><b style="color:#0A0A0A;">IV:</b> {", ".join(ivs)}</div>'
        if meds: rail_html += f'<div style="margin-bottom:4px;"><b style="color:#0A0A0A;">M:</b> {", ".join(meds)}</div>'
        if dvs: rail_html += f'<div><b style="color:#0A0A0A;">DV:</b> {", ".join(dvs)}</div>'
        rail_html += '</div>'

    rail_html += '<h5>Gap Identified</h5>'
    if gap_text:
        rail_html += f'<div class="val">{gap_text[:240]}{"…" if len(gap_text)>240 else ""}</div>'
    else:
        rail_html += '<div class="empty">Complete Stage 03</div>'

    rail_html += '<h5>Method</h5>'
    if method_text:
        rail_html += f'<div class="val">{method_text[:160]}{"…" if len(method_text)>160 else ""}</div>'
    else:
        rail_html += '<div class="empty">Complete Stage 04</div>'

    if score:
        rail_html += f'<h5>Readiness · {journal}</h5>'
        rail_html += f'<div style="font-family:var(--serif);font-size:32px;color:#0A0A0A;font-weight:500;letter-spacing:-0.02em;">{score}<span style="font-size:16px;color:#787776;">/10</span></div>'

    rail_html += '</div>'
    st.markdown(rail_html, unsafe_allow_html=True)

# -------- MAIN COLUMN --------
with main_col:
    cur_idx = st.session_state.current_stage_idx
    cur_key, cur_label, cur_num = STAGES[cur_idx]

    # Page header
    st.markdown(f'''
    <div class="eagle-page-head">
      <div>
        <div class="eagle-kicker">Stage {cur_num}</div>
        <h1 style="margin:0;">{cur_label}</h1>
      </div>
    </div>
    ''', unsafe_allow_html=True)

    # ===================================================================
    # STAGE 1 — PROBLEMATIZE
    # ===================================================================
    if cur_key == "stage_1_problematization":
        s1 = canvas["stage_1_problematization"]
        st.markdown('<p class="eagle-lede">Start with the idea you\'re chasing. One or two sentences — don\'t polish it. Eagle\'s Problematizer will push you with Socratic questions until it\'s sharp.</p>', unsafe_allow_html=True)
        st.markdown("<br/>", unsafe_allow_html=True)

        fuzzy = st.text_area("Fuzzy idea", value=s1["fuzzy_idea"], height=90,
                             placeholder="e.g. I want to study whether AR apps make luxury shoppers feel more confident — and whether that confidence backfires after purchase.")
        if fuzzy != s1["fuzzy_idea"]:
            s1["fuzzy_idea"] = fuzzy
            db_save(canvas)

        # Socratic exchange
        if s1["socratic_exchange"]:
            st.markdown('<div class="eagle-kicker" style="margin-top:24px;">Socratic exchange</div>', unsafe_allow_html=True)
            for i, ex in enumerate(s1["socratic_exchange"]):
                st.markdown(f'''
                <div style="margin-bottom:18px;padding:14px 16px;border-left:2px solid var(--accent);background:var(--bg-elevated);">
                  <div style="font-family:var(--mono);font-size:10px;color:#B4B3B0;letter-spacing:0.1em;">Q{i+1:02d}</div>
                  <div style="font-family:var(--serif);font-size:15px;font-style:italic;color:#0A0A0A;margin-top:2px;">{ex["q"]}</div>
                </div>
                ''', unsafe_allow_html=True)
                if ex["a"]:
                    st.markdown(f'<div style="margin:-10px 0 20px 18px;color:#3F3F3E;font-size:14px;line-height:1.5;">{ex["a"]}</div>', unsafe_allow_html=True)

        # Next question / answer
        if s1["fuzzy_idea"] and len(s1["socratic_exchange"]) < 6 and not s1["candidate_rqs"]:
            needs_answer = s1["socratic_exchange"] and not s1["socratic_exchange"][-1]["a"]
            if needs_answer:
                ans = st.text_area("Your answer", key=f"ans_{len(s1['socratic_exchange'])}", height=90)
                ca, cb = st.columns([1, 4])
                with ca:
                    if st.button("Submit", type="primary"):
                        s1["socratic_exchange"][-1]["a"] = ans
                        db_save(canvas)
                        st.rerun()
            else:
                if st.button("Ask next question", type="primary"):
                    history = f"Fuzzy idea: {s1['fuzzy_idea']}\n\nExchange so far:\n"
                    for ex in s1["socratic_exchange"]:
                        history += f"Q: {ex['q']}\nA: {ex['a']}\n"
                    with st.spinner("Thinking…"):
                        q = claude_call(SYS_PROBLEMATIZER_QUESTION, history).strip().strip('"').strip()
                    s1["socratic_exchange"].append({"q": q, "a": ""})
                    db_save(canvas)
                    st.rerun()

        # Finalize RQs
        if (len(s1["socratic_exchange"]) >= 3
            and all(ex["a"] for ex in s1["socratic_exchange"])
            and not s1["candidate_rqs"]):
            st.markdown("<br/>", unsafe_allow_html=True)
            if st.button("Generate candidate research questions →", type="primary"):
                history = f"Fuzzy idea: {s1['fuzzy_idea']}\n\nExchange:\n"
                for ex in s1["socratic_exchange"]:
                    history += f"Q: {ex['q']}\nA: {ex['a']}\n"
                with st.spinner("Drafting candidates…"):
                    out = claude_json(SYS_PROBLEMATIZER_FINALIZE, history)
                s1["candidate_rqs"] = out["candidate_rqs"]
                db_save(canvas)
                st.rerun()

        if s1["candidate_rqs"]:
            st.markdown('<div class="eagle-kicker" style="margin-top:24px;">Candidate research questions</div>', unsafe_allow_html=True)
            for i, cand in enumerate(s1["candidate_rqs"]):
                selected = s1.get("selected_rq") == cand["rq"]
                border = "var(--accent)" if selected else "var(--rule)"
                bg = "var(--accent-soft)" if selected else "var(--bg-elevated)"
                st.markdown(f'''
                <div style="padding:18px 20px;border:1px solid {border};background:{bg};border-radius:4px;margin-bottom:10px;">
                  <div style="font-family:var(--mono);font-size:10px;color:#B4B3B0;letter-spacing:0.1em;">OPTION {i+1:02d}</div>
                  <div style="font-family:var(--serif);font-size:17px;color:#0A0A0A;margin-top:4px;line-height:1.35;">{cand["rq"]}</div>
                  <div style="font-size:12px;color:#787776;margin-top:8px;font-style:italic;">Trade-off: {cand["tradeoff"]}</div>
                </div>
                ''', unsafe_allow_html=True)
            choice = st.radio("Select your RQ",
                              options=[c["rq"] for c in s1["candidate_rqs"]],
                              index=None, label_visibility="collapsed")
            if choice and choice != s1["selected_rq"]:
                s1["selected_rq"] = choice
                db_save(canvas)
                st.success("Research question locked. Continue to Stage 02 — Theorize.")

    # ===================================================================
    # STAGE 2 — THEORIZE
    # ===================================================================
    elif cur_key == "stage_2_theorization":
        s1 = canvas["stage_1_problematization"]
        s2 = canvas["stage_2_theorization"]

        if not s1["selected_rq"]:
            st.info("Select a research question in Stage 01 first.")
        else:
            st.markdown(f'''
            <div class="eagle-card" style="border-left:3px solid var(--accent);">
              <div class="eagle-kicker">Research question</div>
              <div style="font-family:var(--serif);font-size:18px;line-height:1.4;color:#0A0A0A;">{s1["selected_rq"]}</div>
            </div>
            ''', unsafe_allow_html=True)

            if not s2["candidate_theories"]:
                if st.button("Propose theoretical lenses & construct map →", type="primary"):
                    with st.spinner("Theorist working…"):
                        out = claude_json(SYS_THEORIST, f"RQ: {s1['selected_rq']}")
                    s2["candidate_theories"] = out.get("candidate_theories", [])
                    s2["selected_theory"] = out.get("recommended_theory", "")
                    s2["recommendation_rationale"] = out.get("recommendation_rationale","")
                    s2["construct_map"] = out.get("construct_map", {})
                    s2["hypotheses"] = out.get("hypotheses", [])
                    db_save(canvas)
                    st.rerun()
            else:
                # candidate theories
                st.markdown('<div class="eagle-kicker" style="margin-top:24px;">Candidate theories</div>', unsafe_allow_html=True)
                for t in s2["candidate_theories"]:
                    is_selected = t["name"] == s2["selected_theory"]
                    marker = '<span class="chip accent">recommended</span>' if is_selected else ""
                    with st.expander(f"{t['name']}  {'✦' if is_selected else ''}", expanded=is_selected):
                        st.markdown(f"**Core claim.** {t.get('core_claim','')}")
                        st.markdown(f"**Strengths.** {t.get('strengths','')}")
                        st.markdown(f"**Weaknesses.** {t.get('weaknesses','')}")
                        refs = t.get("seminal_refs", [])
                        if refs:
                            st.markdown(f"**Seminal references.** {'; '.join(refs)}")

                chosen = st.text_input("Selected theory (edit if needed)",
                                       value=s2["selected_theory"])
                if chosen != s2["selected_theory"]:
                    s2["selected_theory"] = chosen
                    db_save(canvas)

                if s2.get("recommendation_rationale"):
                    st.markdown(f'<div style="font-style:italic;color:#3F3F3E;font-size:13px;margin-top:-8px;margin-bottom:16px;">{s2["recommendation_rationale"]}</div>', unsafe_allow_html=True)

                # ============ HYPOTHESIS PATH DIAGRAM ============
                if s2.get("construct_map"):
                    st.markdown('<div class="eagle-kicker" style="margin-top:28px;">Hypothesis path diagram</div>', unsafe_allow_html=True)
                    st.markdown('<p style="font-size:13px;color:#787776;margin-bottom:12px;">Auto-generated from the construct map. Hypothesis labels (H1, H2…) mark each path.</p>', unsafe_allow_html=True)
                    svg = render_path_diagram(s2["construct_map"], s2.get("hypotheses",[]))
                    components.html(f'<div style="width:100%;">{svg}</div>',
                                    height=max(520, 140 + max(
                                        len(s2["construct_map"].get("focal_iv",[])),
                                        len(s2["construct_map"].get("focal_dv",[])),
                                        len(s2["construct_map"].get("mediators",[])), 1) * 110), scrolling=False)

                # hypotheses list
                if s2.get("hypotheses"):
                    st.markdown('<div class="eagle-kicker" style="margin-top:24px;">Hypotheses</div>', unsafe_allow_html=True)
                    for h in s2["hypotheses"]:
                        htype = h.get("type","direct")
                        st.markdown(f'''
                        <div style="padding:12px 16px;border-left:2px solid var(--ink);background:var(--bg-elevated);margin-bottom:8px;border-radius:0 4px 4px 0;">
                          <span style="font-family:var(--mono);font-size:11px;color:var(--accent);font-weight:600;">{h.get('id','H')}</span>
                          <span style="font-family:var(--mono);font-size:10px;color:#B4B3B0;margin-left:8px;">{htype}</span>
                          <div style="font-size:14px;color:#0A0A0A;margin-top:4px;line-height:1.5;">{h.get('text','')}</div>
                        </div>
                        ''', unsafe_allow_html=True)

                # construct map raw (for inspection)
                with st.expander("Construct map (raw)"):
                    st.json(s2["construct_map"])

                cA, cB = st.columns([1, 4])
                with cA:
                    if st.button("Regenerate", type="secondary"):
                        s2["candidate_theories"] = []
                        s2["construct_map"] = {}
                        s2["hypotheses"] = []
                        db_save(canvas)
                        st.rerun()

    # ===================================================================
    # STAGE 3 — LITERATURE
    # ===================================================================
    elif cur_key == "stage_3_literature":
        s2 = canvas["stage_2_theorization"]
        s3 = canvas["stage_3_literature"]

        if not s2.get("construct_map"):
            st.info("Complete Stage 02 first — the Literature Scout needs your construct map.")
        else:
            st.markdown('<p class="eagle-lede">Eagle queries OpenAlex — 240M+ peer-reviewed works — using your constructs and theory. Never fabricates a citation. What you see came back from the API.</p>', unsafe_allow_html=True)

            if not s3["lit_table"]:
                if st.button("Search OpenAlex & build literature review →", type="primary"):
                    context = json.dumps({
                        "theory": s2["selected_theory"],
                        "construct_map": s2["construct_map"],
                        "rq": canvas["stage_1_problematization"]["selected_rq"],
                    })
                    with st.spinner("Drafting search queries…"):
                        q_out = claude_json(SYS_LITSCOUT_QUERIES, context, model=MODEL_LIGHT)
                    queries = q_out.get("queries", [])
                    s3["search_queries_used"] = queries

                    all_papers = []
                    seen = set()
                    progress = st.progress(0.0, "Querying OpenAlex…")
                    for i, q in enumerate(queries):
                        papers = openalex_search(q, per_page=20)
                        for p in papers:
                            if p["openalex_id"] not in seen:
                                seen.add(p["openalex_id"])
                                all_papers.append(p)
                        progress.progress((i+1)/len(queries), f"Query: {q}")
                    progress.empty()

                    st.info(f"Retrieved {len(all_papers)} unique papers. Synthesizing…")
                    all_papers.sort(key=lambda p: p.get("cited_by",0), reverse=True)
                    blob = json.dumps(all_papers[:40])
                    synth_input = (
                        f"RQ: {canvas['stage_1_problematization']['selected_rq']}\n"
                        f"Theory: {s2['selected_theory']}\n"
                        f"Construct map: {json.dumps(s2['construct_map'])}\n\n"
                        f"Papers from OpenAlex:\n{blob}"
                    )
                    with st.spinner("Building lit table and synthesis…"):
                        out = claude_json(SYS_LITSCOUT_SYNTHESIZE, synth_input)
                    s3["lit_table"] = out.get("lit_table", [])
                    s3["synthesis"] = out.get("synthesis", "")
                    s3["identified_gap"] = out.get("identified_gap", "")
                    db_save(canvas)
                    st.rerun()
            else:
                # queries used
                st.markdown('<div class="eagle-kicker">Queries used</div>', unsafe_allow_html=True)
                chips = "".join(f'<span class="chip">{q}</span>' for q in s3["search_queries_used"])
                st.markdown(f'<div style="margin-bottom:20px;">{chips}</div>', unsafe_allow_html=True)

                # gap
                if s3.get("identified_gap"):
                    st.markdown(f'''
                    <div class="eagle-card" style="border-left:3px solid var(--accent);background:var(--accent-soft);border-color:var(--accent);">
                      <div class="eagle-kicker" style="color:var(--accent);">Gap identified</div>
                      <div style="font-family:var(--serif);font-size:16px;line-height:1.45;color:#0A0A0A;">{s3["identified_gap"]}</div>
                    </div>
                    ''', unsafe_allow_html=True)

                # synthesis
                if s3.get("synthesis"):
                    st.markdown('<div class="eagle-kicker" style="margin-top:24px;">Synthesis</div>', unsafe_allow_html=True)
                    st.markdown(f'<div style="font-size:14px;line-height:1.65;color:#3F3F3E;">{s3["synthesis"]}</div>', unsafe_allow_html=True)

                # lit table as editorial cards
                st.markdown('<div class="eagle-kicker" style="margin-top:28px;">Literature</div>', unsafe_allow_html=True)
                for p in s3["lit_table"]:
                    doi = p.get("doi","")
                    doi_link = f'<a href="{doi}" target="_blank" style="color:#787776;font-family:var(--mono);font-size:11px;">↗ DOI</a>' if doi else ''
                    st.markdown(f'''
                    <div style="padding:16px 18px;border:1px solid var(--rule);background:var(--bg-elevated);border-radius:4px;margin-bottom:10px;">
                      <div style="display:flex;justify-content:space-between;align-items:baseline;">
                        <div style="font-family:var(--serif);font-size:15px;color:#0A0A0A;font-weight:500;">{p.get('authors','')} <span style="color:#787776;font-weight:400;">({p.get('year','')})</span></div>
                        {doi_link}
                      </div>
                      <div style="font-size:12px;color:#787776;font-style:italic;margin-top:2px;">{p.get('venue','')}</div>
                      <div style="margin-top:10px;font-size:13px;line-height:1.5;color:#3F3F3E;"><b style="color:#0A0A0A;">Findings.</b> {p.get('findings','')}</div>
                      <div style="margin-top:6px;font-size:13px;line-height:1.5;color:#3F3F3E;"><b style="color:#0A0A0A;">Gap.</b> {p.get('gap_identified','')}</div>
                      <div style="margin-top:10px;font-family:var(--mono);font-size:11px;color:#B4B3B0;">
                        {p.get('theory','')} · {p.get('method','')} · {p.get('context','')}
                      </div>
                    </div>
                    ''', unsafe_allow_html=True)

                if st.button("Re-run literature search", type="secondary"):
                    s3["lit_table"] = []
                    s3["synthesis"] = ""
                    db_save(canvas)
                    st.rerun()

    # ===================================================================
    # STAGE 4 — METHOD + POWER CALCULATOR
    # ===================================================================
    elif cur_key == "stage_4_method":
        s3 = canvas["stage_3_literature"]
        s4 = canvas["stage_4_method"]

        if not s3.get("lit_table"):
            st.info("Complete Stage 03 first.")
        else:
            # approach selector
            st.markdown('<div class="eagle-kicker">Research approach</div>', unsafe_allow_html=True)
            approach = st.radio(
                "Approach", options=["quantitative","qualitative"],
                index=0 if s4.get("research_approach","quantitative")=="quantitative" else 1,
                horizontal=True, label_visibility="collapsed"
            )
            if approach != s4.get("research_approach"):
                s4["research_approach"] = approach
                s4["instrument"] = {"blocks":[],"attention_checks":[],"demographics":[]}
                s4["interview_protocol"] = {"sections":[],"probes":[],"sampling_strategy":""}
                db_save(canvas)
                st.rerun()

            # ============ POWER CALCULATOR (quantitative only) ============
            if approach == "quantitative":
                st.markdown('<div class="eagle-kicker" style="margin-top:24px;">Interactive power analysis</div>', unsafe_allow_html=True)
                st.markdown('<p style="font-size:13px;color:#787776;margin-bottom:12px;">Drag the sliders. Required sample size updates live. G*Power-style F-test calculation.</p>', unsafe_allow_html=True)
                components.html(POWER_WIDGET_HTML, height=400, scrolling=False)

                # simple persistence for computed N — default based on current values
                with st.expander("Save current power calculation to canvas"):
                    c1, c2, c3, c4, c5 = st.columns(5)
                    with c1: f2 = st.number_input("f²", value=float(s4["power_calc"].get("effect_f2",0.15)), step=0.01, format="%.2f")
                    with c2: al = st.number_input("α", value=float(s4["power_calc"].get("alpha",0.05)), step=0.01, format="%.2f")
                    with c3: pw = st.number_input("Power", value=float(s4["power_calc"].get("power",0.80)), step=0.01, format="%.2f")
                    with c4: k = st.number_input("Predictors", value=int(s4["power_calc"].get("n_predictors",5)), step=1)
                    with c5:
                        # recompute n using the same python formula as the JS
                        def python_power_n(f2, alpha, power, k):
                            Ltable = {1:7.85,2:9.64,3:10.90,4:11.94,5:12.83,6:13.62,7:14.35,8:15.02,10:16.25,15:18.88,20:21.20}
                            pFac = {0.70:0.85,0.75:0.92,0.80:1.00,0.85:1.10,0.90:1.22,0.95:1.40,0.99:1.75}
                            aFac = {0.01:1.50,0.02:1.33,0.05:1.00,0.10:0.78}
                            def interp(tbl, val):
                                keys = sorted(tbl.keys())
                                if val <= keys[0]: return tbl[keys[0]]
                                if val >= keys[-1]: return tbl[keys[-1]]
                                for i in range(len(keys)-1):
                                    if keys[i] <= val <= keys[i+1]:
                                        f = (val-keys[i])/(keys[i+1]-keys[i])
                                        return tbl[keys[i]] + f*(tbl[keys[i+1]]-tbl[keys[i]])
                                return 1.0
                            L = interp(Ltable, k) * interp(pFac, power) * interp(aFac, alpha)
                            n = int(L / f2) + k + 1
                            return max(n, k + 3)
                        n_comp = python_power_n(f2, al, pw, k)
                        st.metric("n required", n_comp)
                    if st.button("Save to canvas", type="secondary"):
                        s4["power_calc"] = {"effect_f2":f2,"alpha":al,"power":pw,"n_predictors":k,"computed_n":n_comp}
                        if not s4.get("sample_plan"):
                            s4["sample_plan"] = {}
                        s4["sample_plan"]["target_n"] = n_comp
                        s4["sample_plan"]["power_analysis"] = f"G*Power: f²={f2}, α={al}, 1−β={pw}, predictors={k} → n={n_comp}"
                        db_save(canvas)
                        st.success(f"Saved. Target n set to {n_comp}.")

            # ============ METHOD DESIGN ============
            st.markdown("<hr/>", unsafe_allow_html=True)
            has_instrument = bool(s4.get("instrument",{}).get("blocks"))
            has_protocol = bool(s4.get("interview_protocol",{}).get("sections"))

            if not (has_instrument or has_protocol):
                btn_label = "Design survey instrument →" if approach == "quantitative" else "Design interview protocol →"
                if st.button(btn_label, type="primary"):
                    context = json.dumps({
                        "rq": canvas["stage_1_problematization"]["selected_rq"],
                        "theory": canvas["stage_2_theorization"]["selected_theory"],
                        "construct_map": canvas["stage_2_theorization"]["construct_map"],
                        "hypotheses": canvas["stage_2_theorization"]["hypotheses"],
                        "gap": s3["identified_gap"],
                        "target_n": s4.get("power_calc",{}).get("computed_n",""),
                    })
                    sys_prompt = SYS_METHOD_QUANT if approach == "quantitative" else SYS_METHOD_QUAL
                    with st.spinner("Designing method and instrument…"):
                        out = claude_json(sys_prompt, context)
                    s4["recommended_method"] = out.get("recommended_method","")
                    s4["method_rationale"] = out.get("method_rationale","")
                    s4["sample_plan"].update(out.get("sample_plan",{}))
                    if approach == "quantitative":
                        s4["instrument"] = out.get("instrument",{})
                    else:
                        s4["interview_protocol"] = out.get("interview_protocol",{})
                    s4["ethics_flags"] = out.get("ethics_flags",[])
                    db_save(canvas)
                    st.rerun()
            else:
                # Show method
                st.markdown('<div class="eagle-kicker">Recommended method</div>', unsafe_allow_html=True)
                st.markdown(f'<div style="font-family:var(--serif);font-size:18px;color:#0A0A0A;line-height:1.4;margin-bottom:8px;">{s4["recommended_method"]}</div>', unsafe_allow_html=True)
                st.markdown(f'<div style="font-size:13px;color:#3F3F3E;line-height:1.6;margin-bottom:20px;">{s4["method_rationale"]}</div>', unsafe_allow_html=True)

                sp = s4.get("sample_plan") or {}
                if sp:
                    c1, c2 = st.columns(2)
                    with c1:
                        if sp.get("target_n"): st.metric("Target N", sp.get("target_n",""))
                    with c2:
                        if sp.get("recruitment"): st.markdown(f'<div style="padding-top:12px;"><b>Recruitment.</b> {sp.get("recruitment","")}</div>', unsafe_allow_html=True)

                # Quantitative instrument
                if has_instrument:
                    st.markdown('<div class="eagle-kicker" style="margin-top:24px;">Survey instrument</div>', unsafe_allow_html=True)
                    for block in s4["instrument"]["blocks"]:
                        validated_chip = '<span class="chip ok">validated</span>' if block.get("validated") else '<span class="chip warn">new — pretest</span>'
                        with st.expander(f"{block['block_name']} — {block.get('construct','')}"):
                            st.markdown(f'<div style="margin-bottom:8px;">{validated_chip} <span class="chip">{block.get("scale_source","")}</span></div>', unsafe_allow_html=True)
                            for item in block.get("items",[]):
                                tag = " <span class='chip' style='font-size:9px;'>R</span>" if item.get("reverse_coded") else ""
                                st.markdown(f'<div style="padding:6px 0;border-bottom:1px solid var(--rule);font-size:13px;"><span style="font-family:var(--mono);color:#7B1E3A;font-size:11px;">{item.get("id","")}</span>{tag} &nbsp; {item.get("text","")}</div>', unsafe_allow_html=True)

                    if s4["instrument"].get("attention_checks"):
                        st.markdown('<div class="eagle-kicker" style="margin-top:16px;">Attention checks</div>', unsafe_allow_html=True)
                        for ac in s4["instrument"]["attention_checks"]:
                            st.markdown(f'• <span style="font-family:var(--mono);font-size:11px;color:#7B1E3A;">{ac.get("id","")}</span> — {ac.get("text","")}', unsafe_allow_html=True)

                # Qualitative protocol
                if has_protocol:
                    st.markdown('<div class="eagle-kicker" style="margin-top:24px;">Interview protocol</div>', unsafe_allow_html=True)
                    proto = s4["interview_protocol"]
                    if proto.get("opening_question"):
                        st.markdown(f'**Opening.** {proto["opening_question"]}')
                    for sec in proto.get("sections",[]):
                        with st.expander(f"{sec.get('section_name','')}"):
                            st.caption(f"Construct: {sec.get('construct_explored','')}")
                            for q in sec.get("questions",[]):
                                st.markdown(f'**{q.get("id","")}.** {q.get("text","")}')
                                for probe in q.get("probes",[]):
                                    st.markdown(f'&nbsp;&nbsp;&nbsp;&nbsp;↳ {probe}')
                    if proto.get("closing_question"):
                        st.markdown(f'**Closing.** {proto["closing_question"]}')

                if s4.get("ethics_flags"):
                    st.markdown('<div class="eagle-kicker" style="margin-top:24px;">Ethics flags</div>', unsafe_allow_html=True)
                    for f in s4["ethics_flags"]:
                        st.markdown(f'<div style="padding:8px 12px;border-left:2px solid var(--warn);background:var(--bg-elevated);margin-bottom:6px;font-size:13px;">{f}</div>', unsafe_allow_html=True)

                if st.button("Redesign method", type="secondary"):
                    s4["instrument"] = {"blocks":[],"attention_checks":[],"demographics":[]}
                    s4["interview_protocol"] = {"sections":[],"probes":[],"sampling_strategy":""}
                    db_save(canvas)
                    st.rerun()

    # ===================================================================
    # STAGE 5 — REVIEWER
    # ===================================================================
    elif cur_key == "stage_5_review":
        s4 = canvas["stage_4_method"]
        s5 = canvas["stage_5_review"]

        method_ready = bool(s4.get("instrument",{}).get("blocks") or s4.get("interview_protocol",{}).get("sections"))
        if not method_ready:
            st.info("Complete Stage 04 first.")
        else:
            st.markdown('<p class="eagle-lede">A simulated reviewer at your target journal reads the full canvas and identifies the concerns most likely to trigger desk-reject or major revisions — before submission.</p>', unsafe_allow_html=True)

            journals = ["JAMS","JCR","JMR","MISQ","MIT Sloan Mgmt Review",
                        "Information & Management","IJRDM","Journal of Business Research","Other"]
            default_idx = journals.index(s5["target_journal"]) if s5.get("target_journal") in journals else 0
            s5["target_journal"] = st.selectbox("Target journal", journals, index=default_idx)

            if not s5["reviewer_critique"]:
                if st.button(f"Run reviewer critique for {s5['target_journal']} →", type="primary"):
                    review_canvas = {k: v for k, v in canvas.items() if k != "stage_5_review"}
                    review_canvas["target_journal"] = s5["target_journal"]
                    with st.spinner(f"Reviewer reading {s5['target_journal']}-calibrated…"):
                        out = claude_json(SYS_REVIEWER, json.dumps(review_canvas))
                    s5["reviewer_critique"] = out.get("reviewer_critique",[])
                    s5["overall_readiness_score"] = out.get("overall_readiness_score",0)
                    s5["journal_specific_notes"] = out.get("journal_specific_notes","")
                    db_save(canvas)
                    st.rerun()
            else:
                # readiness score big
                score = s5["overall_readiness_score"]
                color = "var(--ok)" if score >= 7 else "var(--warn)" if score >= 5 else "var(--danger)"
                st.markdown(f'''
                <div class="eagle-card" style="display:flex;align-items:baseline;gap:24px;">
                  <div>
                    <div class="eagle-kicker">Readiness · {s5["target_journal"]}</div>
                    <div style="font-family:var(--serif);font-size:72px;font-weight:400;color:{color};line-height:1;letter-spacing:-0.04em;">{score}<span style="font-size:24px;color:#787776;">/10</span></div>
                  </div>
                  <div style="flex:1;font-size:13px;color:#3F3F3E;line-height:1.55;padding-top:20px;">{s5.get("journal_specific_notes","")}</div>
                </div>
                ''', unsafe_allow_html=True)

                # concerns grouped by severity
                by_sev = {"high":[], "medium":[], "low":[]}
                for c in s5["reviewer_critique"]:
                    by_sev.get(c.get("severity","medium"), by_sev["medium"]).append(c)

                for sev, label, color, icon in [("high","High severity","var(--danger)","▲"),
                                                 ("medium","Moderate","var(--warn)","■"),
                                                 ("low","Low severity","var(--ok)","◆")]:
                    if by_sev[sev]:
                        st.markdown(f'<div class="eagle-kicker" style="margin-top:24px;color:{color};">{icon} {label}</div>', unsafe_allow_html=True)
                        for c in by_sev[sev]:
                            st.markdown(f'''
                            <div style="padding:14px 18px;border-left:2px solid {color};background:var(--bg-elevated);margin-bottom:10px;border-radius:0 4px 4px 0;">
                              <div style="display:flex;justify-content:space-between;align-items:baseline;">
                                <span style="font-family:var(--mono);font-size:11px;color:{color};font-weight:600;">{c.get("concern_id","")}</span>
                                <span style="font-family:var(--mono);font-size:10px;color:#B4B3B0;">{c.get("category","")}</span>
                              </div>
                              <div style="font-size:14px;color:#0A0A0A;margin-top:6px;line-height:1.5;"><b>Issue.</b> {c.get("issue","")}</div>
                              <div style="font-size:13px;color:#3F3F3E;margin-top:6px;line-height:1.5;"><b>Fix.</b> {c.get("suggested_fix","")}</div>
                              <div style="font-family:var(--mono);font-size:10px;color:#B4B3B0;margin-top:8px;">→ revise in {c.get("linked_stage","")}</div>
                            </div>
                            ''', unsafe_allow_html=True)

                if st.button("Re-run review", type="secondary"):
                    s5["reviewer_critique"] = []
                    db_save(canvas)
                    st.rerun()

    # ===================================================================
    # STAGE 6 — MANUSCRIPT
    # ===================================================================
    elif cur_key == "stage_6_full_paper":
        s5 = canvas["stage_5_review"]
        s6 = canvas["stage_6_full_paper"]

        if not s5.get("reviewer_critique"):
            st.info("Complete Stage 05 first. The manuscript auto-fixes reviewer concerns.")
        else:
            st.markdown('<p class="eagle-lede">Eagle addresses every reviewer concern and drafts a full manuscript: Introduction, Literature Review, Hypothesis Development, Method, and 4+4 Implications. Export as Word and refine.</p>', unsafe_allow_html=True)

            if not s6.get("generated"):
                if st.button("Generate manuscript →", type="primary"):
                    context = json.dumps({
                        "rq": canvas["stage_1_problematization"]["selected_rq"],
                        "theory": canvas["stage_2_theorization"]["selected_theory"],
                        "construct_map": canvas["stage_2_theorization"]["construct_map"],
                        "hypotheses": canvas["stage_2_theorization"]["hypotheses"],
                        "lit_table": canvas["stage_3_literature"]["lit_table"][:15],
                        "synthesis": canvas["stage_3_literature"]["synthesis"],
                        "gap": canvas["stage_3_literature"]["identified_gap"],
                        "research_approach": canvas["stage_4_method"].get("research_approach","quantitative"),
                        "method": canvas["stage_4_method"]["recommended_method"],
                        "method_rationale": canvas["stage_4_method"]["method_rationale"],
                        "sample_plan": canvas["stage_4_method"]["sample_plan"],
                        "reviewer_critique": s5["reviewer_critique"],
                        "target_journal": s5["target_journal"],
                    })
                    with st.spinner("Writing manuscript (this takes 45-60 seconds)…"):
                        out = claude_json(SYS_FULL_PAPER, context, max_tokens=12000)
                    s6.update({
                        "introduction": out.get("introduction",""),
                        "literature_review": out.get("literature_review",""),
                        "hypotheses_text": out.get("hypotheses_text",""),
                        "method_section": out.get("method_section",""),
                        "theoretical_implications": out.get("theoretical_implications",[]),
                        "practical_implications": out.get("practical_implications",[]),
                        "revised_notes": out.get("revised_notes",""),
                        "generated": True,
                    })
                    db_save(canvas)
                    st.rerun()
            else:
                # Display sections as tabs for readability
                section_tabs = st.tabs(["Introduction","Literature Review","Hypotheses","Method","Implications","How concerns were addressed"])
                with section_tabs[0]:
                    st.markdown(f'<div style="font-family:var(--serif);font-size:15px;line-height:1.7;color:#0A0A0A;">{s6["introduction"].replace(chr(10)+chr(10), "</p><p>")}</div>', unsafe_allow_html=True)
                with section_tabs[1]:
                    st.markdown(f'<div style="font-family:var(--serif);font-size:15px;line-height:1.7;color:#0A0A0A;">{s6["literature_review"].replace(chr(10)+chr(10), "</p><p>")}</div>', unsafe_allow_html=True)
                with section_tabs[2]:
                    st.markdown(f'<div style="font-family:var(--serif);font-size:15px;line-height:1.7;color:#0A0A0A;">{s6["hypotheses_text"].replace(chr(10)+chr(10), "</p><p>")}</div>', unsafe_allow_html=True)
                with section_tabs[3]:
                    st.markdown(f'<div style="font-family:var(--serif);font-size:15px;line-height:1.7;color:#0A0A0A;">{s6["method_section"].replace(chr(10)+chr(10), "</p><p>")}</div>', unsafe_allow_html=True)
                with section_tabs[4]:
                    st.markdown('<div class="eagle-kicker">Theoretical contributions</div>', unsafe_allow_html=True)
                    for imp in s6.get("theoretical_implications",[]):
                        st.markdown(f'**{imp.get("title","")}.** {imp.get("text","")}')
                    st.markdown('<div class="eagle-kicker" style="margin-top:20px;">Practical implications</div>', unsafe_allow_html=True)
                    for imp in s6.get("practical_implications",[]):
                        st.markdown(f'**{imp.get("title","")}.** {imp.get("text","")}')
                with section_tabs[5]:
                    st.markdown(s6.get("revised_notes",""))

                if st.button("Regenerate manuscript", type="secondary"):
                    s6["generated"] = False
                    db_save(canvas)
                    st.rerun()

    # ===================================================================
    # STAGE 7 — RESPONSE LETTER (NEW)
    # ===================================================================
    elif cur_key == "stage_7_response_letter":
        s7 = canvas["stage_7_response_letter"]
        st.markdown('<p class="eagle-lede">Got an R&R? Paste the reviewer comments from your journal (any format — Eagle handles R1/R2/R3, editor letters, numbered lists, or flowing prose). Eagle drafts a point-by-point response letter in journal-ready format.</p>', unsafe_allow_html=True)

        raw = st.text_area(
            "Paste reviewer comments here",
            value=s7.get("reviewer_comments_raw",""),
            height=260,
            placeholder="""Reviewer 1:
1. The theoretical contribution is unclear. The paper reads as a context extension rather than a theoretical advance...
2. The identification strategy for the mediation test is problematic because...

Reviewer 2:
This is an important topic but I have several concerns:
(a) The sample is limited to...
(b) Common method bias is not adequately addressed...

Editor:
Please pay special attention to..."""
        )
        if raw != s7.get("reviewer_comments_raw",""):
            s7["reviewer_comments_raw"] = raw
            db_save(canvas)

        if raw and not s7.get("generated"):
            cA, cB = st.columns([1,4])
            with cA:
                if st.button("Parse & draft response →", type="primary"):
                    # 1) parse
                    with st.spinner("Parsing reviewer comments…"):
                        parsed = claude_json(SYS_PARSE_REVIEWER_COMMENTS, raw, model=MODEL_LIGHT)
                    s7["parsed_comments"] = parsed.get("parsed_comments",[])

                    # 2) draft letter
                    ctx = json.dumps({
                        "canvas": {
                            "title": canvas["title"],
                            "rq": canvas["stage_1_problematization"]["selected_rq"],
                            "theory": canvas["stage_2_theorization"]["selected_theory"],
                            "construct_map": canvas["stage_2_theorization"]["construct_map"],
                            "hypotheses": canvas["stage_2_theorization"]["hypotheses"],
                            "method": canvas["stage_4_method"]["recommended_method"],
                            "target_journal": canvas["stage_5_review"].get("target_journal",""),
                        },
                        "parsed_comments": s7["parsed_comments"],
                        "overall_tone": parsed.get("overall_tone",""),
                        "apparent_decision": parsed.get("apparent_decision",""),
                    })
                    with st.spinner("Drafting point-by-point response letter…"):
                        out = claude_json(SYS_RESPONSE_LETTER, ctx, max_tokens=10000)
                    s7["response_letter"] = out.get("response_letter","")
                    s7["revision_summary"] = out.get("revision_summary","")
                    s7["generated"] = True
                    db_save(canvas)
                    st.rerun()

        # show parsed concerns
        if s7.get("parsed_comments"):
            st.markdown('<div class="eagle-kicker" style="margin-top:24px;">Parsed concerns</div>', unsafe_allow_html=True)
            sources = {}
            for c in s7["parsed_comments"]:
                sources.setdefault(c.get("source","Unknown"), []).append(c)
            for source, items in sources.items():
                st.markdown(f'<div style="font-family:var(--serif);font-size:16px;font-weight:500;margin-top:16px;color:#0A0A0A;">{source}</div>', unsafe_allow_html=True)
                for c in items:
                    sev = c.get("severity_hint","moderate")
                    sev_color = {"major":"var(--danger)","moderate":"var(--warn)","minor":"var(--ok)"}.get(sev,"var(--ink-muted)")
                    st.markdown(f'''
                    <div style="padding:12px 16px;border:1px solid var(--rule);background:var(--bg-elevated);border-radius:4px;margin-bottom:8px;">
                      <div style="display:flex;justify-content:space-between;align-items:baseline;">
                        <span style="font-family:var(--mono);font-size:11px;color:var(--accent);font-weight:600;">{c.get("id","")}</span>
                        <span><span class="chip">{c.get("category","")}</span><span style="font-family:var(--mono);font-size:10px;color:{sev_color};font-weight:600;margin-left:6px;">{sev}</span></span>
                      </div>
                      <div style="font-size:13px;color:#3F3F3E;margin-top:6px;line-height:1.5;">{c.get("comment_text","")}</div>
                    </div>
                    ''', unsafe_allow_html=True)

        # show drafted letter
        if s7.get("generated"):
            st.markdown("<hr/>", unsafe_allow_html=True)
            if s7.get("revision_summary"):
                st.markdown('<div class="eagle-kicker">Revision summary (for cover letter)</div>', unsafe_allow_html=True)
                st.markdown(f'<div style="padding:18px 20px;border-left:3px solid var(--accent);background:var(--accent-soft);font-family:var(--serif);font-size:14px;line-height:1.6;color:#0A0A0A;margin-bottom:20px;">{s7["revision_summary"]}</div>', unsafe_allow_html=True)

            st.markdown('<div class="eagle-kicker">Point-by-point response</div>', unsafe_allow_html=True)
            st.markdown(f'<div style="padding:24px;background:var(--bg-elevated);border:1px solid var(--rule);border-radius:4px;font-family:var(--serif);font-size:14px;line-height:1.7;color:#0A0A0A;">{s7["response_letter"].replace(chr(10), "<br/>")}</div>', unsafe_allow_html=True)

            # Dedicated response letter download
            st.markdown("<br/>", unsafe_allow_html=True)
            try:
                rl_bytes = export_response_letter_docx(canvas)
                cA, cB = st.columns([1,4])
                with cA:
                    st.download_button(
                        "Download response letter (.docx)",
                        data=rl_bytes,
                        file_name=f"response_letter_{canvas['session_id'][:8]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                with cB:
                    if st.button("Redraft response letter", type="secondary"):
                        s7["generated"] = False
                        s7["response_letter"] = ""
                        db_save(canvas)
                        st.rerun()
            except Exception as e:
                st.error(f"Export error: {e}")
